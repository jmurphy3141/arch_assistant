"""
Render generated JEP Markdown into the approved C3E Word template.
"""
from __future__ import annotations

import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


TEMPLATE_PATH = (
    Path(__file__).resolve().parent
    / "standards"
    / "templates"
    / "jep"
    / "c3e_jep_template.docx"
)
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def render_jep_docx(
    markdown: str,
    *,
    customer_name: str = "",
    template_path: Path | str = TEMPLATE_PATH,
) -> bytes:
    """
    Render JEP Markdown into DOCX bytes using the C3E template as style source.
    """
    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"JEP DOCX template not found: {template}")

    doc = Document(str(template))
    _clear_document_body(doc)
    if customer_name:
        doc.core_properties.subject = customer_name

    blocks = list(_iter_blocks(markdown))
    if not blocks:
        _add_paragraph(doc, "Joint Execution Plan", "Title")
    for block_type, value in blocks:
        if block_type == "heading":
            level, text = value
            style = "Title" if level == 1 else f"Heading {min(level - 1, 4)}"
            _add_paragraph(doc, text, style)
        elif block_type == "bullet":
            _add_paragraph(doc, value, "List Bullet")
        elif block_type == "number":
            _add_paragraph(doc, value, "List Number")
        elif block_type == "table":
            _add_table(doc, value)
        elif block_type == "paragraph":
            _add_paragraph(doc, value, "Normal")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _clear_document_body(doc: DocumentObject) -> None:
    body = doc._body._element
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = deepcopy(child)
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def _iter_blocks(markdown: str) -> Iterable[tuple[str, object]]:
    lines = str(markdown or "").splitlines()
    idx = 0
    paragraph: list[str] = []

    def flush_paragraph() -> tuple[str, str] | None:
        nonlocal paragraph
        if not paragraph:
            return None
        text = " ".join(part.strip() for part in paragraph if part.strip()).strip()
        paragraph = []
        if not text:
            return None
        return ("paragraph", _clean_inline(text))

    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()
        if not line:
            block = flush_paragraph()
            if block:
                yield block
            idx += 1
            continue
        if re.fullmatch(r"-{3,}", line):
            block = flush_paragraph()
            if block:
                yield block
            idx += 1
            continue
        if _is_table_line(line):
            block = flush_paragraph()
            if block:
                yield block
            table_lines: list[str] = []
            while idx < len(lines) and _is_table_line(lines[idx].strip()):
                table_lines.append(lines[idx].strip())
                idx += 1
            rows = _parse_table(table_lines)
            if rows:
                yield ("table", rows)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            block = flush_paragraph()
            if block:
                yield block
            yield ("heading", (len(heading.group(1)), _clean_inline(heading.group(2))))
            idx += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            block = flush_paragraph()
            if block:
                yield block
            yield ("bullet", _clean_inline(bullet.group(1)))
            idx += 1
            continue
        numbered = re.match(r"^\d+[.)]\s+(.+)$", line)
        if numbered:
            block = flush_paragraph()
            if block:
                yield block
            yield ("number", _clean_inline(numbered.group(1)))
            idx += 1
            continue
        paragraph.append(line)
        idx += 1

    block = flush_paragraph()
    if block:
        yield block


def _is_table_line(line: str) -> bool:
    return line.startswith("|") and line.count("|") >= 2


def _parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if not cells:
            continue
        if all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append([_clean_inline(cell) for cell in cells])
    return rows


def _add_paragraph(doc: DocumentObject, text: str, style_name: str) -> Paragraph:
    try:
        return doc.add_paragraph(text, style=style_name)
    except Exception:
        return doc.add_paragraph(text)


def _add_table(doc: DocumentObject, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=width)
    for style_name in ("Table Grid", "Normal Table"):
        try:
            table.style = style_name
            break
        except Exception:
            pass
    for row_idx, row_values in enumerate(rows):
        row = table.add_row()
        for col_idx in range(width):
            cell = row.cells[col_idx]
            cell.text = row_values[col_idx] if col_idx < len(row_values) else ""
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _clean_inline(text: str) -> str:
    cleaned = str(text or "").strip()
    cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__([^_]+)__", r"\1", cleaned)
    cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
    cleaned = re.sub(r"_([^_]+)_", r"\1", cleaned)
    return cleaned.strip()
