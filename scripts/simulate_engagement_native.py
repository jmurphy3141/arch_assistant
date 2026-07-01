#!/usr/bin/env python3
"""Run the live Northwind Health engagement against an isolated native stack.

Isolated bring-up used for this qualification (ports are intentionally separate
from production):

1. Copy the current source tree to ``/tmp/archie-native-sim`` without runtime
   output or Git metadata.
2. In that copy only, set ``orchestrator.agent_mode`` to ``native`` and point
   ``sub_agents`` at 127.0.0.1 ports 18082-18087. Keep the committed
   ``config.yaml`` default at ``forge``.
3. Start these services from the isolated copy:

   python3.11 -m uvicorn sub_agents.diagram.server:app --port 18082
   python3.11 -m uvicorn sub_agents.bom.server:app --port 18083
   python3.11 -m uvicorn sub_agents.pov.server:app --port 18084
   python3.11 -m uvicorn sub_agents.jep.server:app --port 18085
   python3.11 -m uvicorn sub_agents.waf.server:app --port 18086
   python3.11 -m uvicorn sub_agents.terraform.server:app --port 18087
   A2A_BASE_URL=http://127.0.0.1:18080 python3.11 -m uvicorn \
       drawing_agent_server:app --host 127.0.0.1 --port 18080

4. Run this script from the source checkout:

   python3.11 scripts/simulate_engagement_native.py \
       --base-url http://127.0.0.1:18080

The run uses a timestamped customer_id, so its real OCI Object Storage evidence
cannot collide with a production engagement. Model behavior is recorded as-is;
the harness never retries a failed intent with command-like phrasing.
"""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
from io import BytesIO
import json
from pathlib import Path
import re
import sys
import time
from typing import Any
from urllib import error, request
import uuid
from xml.etree import ElementTree

import openpyxl
import yaml
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from agent.object_store_oci import OciObjectStore


DEFAULT_REPORT = ROOT / "docs" / "native-engagement-sim.json"
LOOKUP_TOOLS = {"get_document", "get_summary", "list_documents"}
NOTES = """Northwind call scratchpad — rough, please clean up

Member portal is .NET and still in their data center. Healthcare/member claims,
so HIPAA and audit evidence will matter. Priya Shah is the Northwind application
owner; Marcus Lee owns security. Oracle SA partner is Jordan Kim.

For a proof point they mentioned 200 concurrent portal users, claims lookup p95
under 600 ms, and service recovery within 45 minutes. They can give us 20 business
days for a POC. Nobody picked an OCI region or final topology on this call. No
budget, production SLA, server sizing, or storage quantity was agreed.
"""

TURNS: tuple[dict[str, Any], ...] = (
    {
        "id": 1,
        "session_id": "m1",
        "message": (
            "Just wrapped a first call with Northwind Health. They run a .NET member "
            "portal and are kicking around a move to OCI. Honestly that's about all I got today."
        ),
        "kind": "conversation",
    },
    {
        "id": 2,
        "session_id": "m1",
        "message": "What would you want to find out before we meet again?",
        "kind": "conversation",
    },
    {
        "id": 3,
        "session_id": "m1",
        "message": "Have we written anything up for them yet?",
        "kind": "lookup_empty",
    },
    {
        "id": 4,
        "session_id": "m1",
        "message": "Dropped my call notes in — can you pull out what matters?",
        "kind": "notes",
        "upload_notes": True,
    },
    {
        "id": 5,
        "session_id": "m2",
        "message": (
            "Had our second call. Three tiers — IIS web, a claims app, an Oracle database. "
            "HIPAA is non-negotiable, and they're leaning toward us-chicago-1."
        ),
        "kind": "conversation",
    },
    {
        "id": 6,
        "session_id": "m2",
        "message": "Remind me everything we know about them so far.",
        "kind": "memory",
    },
    {
        "id": 7,
        "session_id": "m2",
        "message": "So what's your gut on the architecture?",
        "kind": "conversation",
    },
    {
        "id": 8,
        "session_id": "m2",
        "message": "I think we're ready — can you pull together a POV for them?",
        "kind": "deliverable",
        "tool": "generate_pov",
        "artifact_type": "pov",
    },
    {
        "id": 9,
        "session_id": "m3",
        "message": (
            "They can give us 20 business days. Jordan Kim owns the Oracle side and "
            "Priya Shah owns it for Northwind; each can commit 8 hours a week. Success "
            "means 200 concurrent users, claims lookup p95 under 600 ms, and recovery "
            "inside 45 minutes. Keep IIS, the claims app, Oracle, HA, security, logging, "
            "and monitoring in scope; production cutover, DR, and load beyond 200 users "
            "are out."
        ),
        "kind": "logistics",
    },
    {
        "id": 10,
        "session_id": "m3",
        "message": "They liked the POV. What POC could we run to prove it out?",
        "kind": "deliverable",
        "tool": "generate_poc_plan",
        "artifact_type": "poc_plan",
        "required_inputs_supplied": True,
    },
    {
        "id": 11,
        "session_id": "m3",
        "message": "Let's go with the second option.",
        "kind": "selection",
    },
    {
        "id": 12,
        "session_id": "m3",
        "message": "Great — can you get the architecture diagram going for that?",
        "kind": "deliverable",
        "tool": "generate_diagram",
        "artifact_type": "diagram",
    },
    {
        "id": 13,
        "session_id": "m3",
        "message": (
            "We'll need a BOM too — figure a couple of web boxes, a few app servers, "
            "an HA Oracle database, and some file storage."
        ),
        "kind": "deliverable",
        "tool": "generate_bom",
        "artifact_type": "bom",
    },
    {
        "id": 14,
        "session_id": "m3",
        "message": "Did that BOM slip in any Gen AI token costs?",
        "kind": "lookup_bom",
    },
    {
        "id": 15,
        "session_id": "m3",
        "message": "Last thing — put the JEP together for it.",
        "kind": "deliverable",
        "tool": "generate_jep",
        "artifact_type": "jep",
        "required_inputs_supplied": True,
    },
    {
        "id": 16,
        "session_id": "wrap",
        "message": "Remind me what we've actually produced for Northwind so far.",
        "kind": "lookup_artifacts",
    },
)


def _post_json(
    base_url: str,
    path: str,
    payload: dict[str, Any],
    timeout: float,
) -> tuple[dict[str, Any], float, int]:
    started = time.monotonic()
    req = request.Request(
        base_url.rstrip("/") + path,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with request.urlopen(req, timeout=timeout) as response:
            status = response.status
            body = json.loads(response.read().decode("utf-8"))
    except error.HTTPError as exc:
        status = exc.code
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            body = json.loads(raw)
        except json.JSONDecodeError:
            body = {"detail": raw}
    return body, time.monotonic() - started, status


def _upload_note(
    base_url: str,
    customer_id: str,
    timeout: float,
) -> tuple[dict[str, Any], float, int]:
    boundary = "----archie-native-sim-" + uuid.uuid4().hex
    chunks = [
        (
            f"--{boundary}\r\n"
            'Content-Disposition: form-data; name="customer_id"\r\n\r\n'
            f"{customer_id}\r\n"
        ).encode(),
        (
            f"--{boundary}\r\n"
            'Content-Disposition: form-data; name="note_name"\r\n\r\n'
            "northwind-call-notes.txt\r\n"
        ).encode(),
        (
            f"--{boundary}\r\n"
            'Content-Disposition: form-data; name="file"; filename="northwind-call-notes.txt"\r\n'
            "Content-Type: text/plain\r\n\r\n"
        ).encode()
        + NOTES.encode("utf-8")
        + b"\r\n",
        f"--{boundary}--\r\n".encode(),
    ]
    req = request.Request(
        base_url.rstrip("/") + "/api/notes/upload",
        data=b"".join(chunks),
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    started = time.monotonic()
    try:
        with request.urlopen(req, timeout=timeout) as response:
            status = response.status
            body = json.loads(response.read().decode("utf-8"))
    except error.HTTPError as exc:
        status = exc.code
        raw = exc.read().decode("utf-8", errors="replace")
        try:
            body = json.loads(raw)
        except json.JSONDecodeError:
            body = {"detail": raw}
    return body, time.monotonic() - started, status


def _store_from_config() -> OciObjectStore:
    config = yaml.safe_load((ROOT / "config.yaml").read_text(encoding="utf-8")) or {}
    persistence = config.get("persistence") or {}
    return OciObjectStore(
        region=str(persistence.get("region") or config.get("region") or "us-chicago-1"),
        namespace=str(persistence.get("namespace") or ""),
        bucket_name=str(persistence.get("bucket_name") or ""),
    )


def _tool_calls(body: dict[str, Any]) -> list[dict[str, Any]]:
    calls = body.get("tool_calls") or []
    return [call for call in calls if isinstance(call, dict)]


def _tool_names(body: dict[str, Any]) -> list[str]:
    return [str(call.get("tool") or "") for call in _tool_calls(body)]


def _compact_tool_calls(body: dict[str, Any]) -> list[dict[str, Any]]:
    return [
        {
            "tool": str(call.get("tool") or ""),
            "args": call.get("args") if isinstance(call.get("args"), dict) else {},
            "artifact_key": str(call.get("artifact_key") or ""),
            "result_status": str(call.get("result_status") or ""),
            "result_summary": str(call.get("result_summary") or ""),
        }
        for call in _tool_calls(body)
    ]


def _call_for(body: dict[str, Any], tool_name: str) -> dict[str, Any]:
    return next(
        (call for call in _tool_calls(body) if str(call.get("tool") or "") == tool_name),
        {},
    )


def _manifest_download(body: dict[str, Any], artifact_type: str) -> dict[str, Any]:
    downloads = (body.get("artifact_manifest") or {}).get("downloads") or []
    return next(
        (
            item
            for item in downloads
            if isinstance(item, dict) and str(item.get("type") or "") == artifact_type
        ),
        {},
    )


def _artifact_key(body: dict[str, Any], tool_name: str, artifact_type: str) -> str:
    manifest = _manifest_download(body, artifact_type)
    if manifest.get("key"):
        return str(manifest["key"])
    call = _call_for(body, tool_name)
    if call.get("artifact_key"):
        return str(call["artifact_key"])
    data = call.get("result_data") if isinstance(call.get("result_data"), dict) else {}
    for key_name in (
        "docx_key",
        "xlsx_artifact_key",
        "drawio_key",
        "object_key",
        "artifact_key",
    ):
        if data.get(key_name):
            return str(data[key_name])
    return ""


def _download_manifest_artifact(
    base_url: str,
    body: dict[str, Any],
    artifact_type: str,
    timeout: float,
) -> tuple[bool, str]:
    item = _manifest_download(body, artifact_type)
    url = str(item.get("download_url") or "")
    if not url:
        return False, f"no {artifact_type} download in artifact_manifest"
    try:
        with request.urlopen(base_url.rstrip("/") + url, timeout=timeout) as response:
            payload = response.read()
        return bool(payload), f"downloaded {len(payload)} bytes from {url}"
    except Exception as exc:
        return False, f"artifact download failed: {exc}"


def _artifact_validation_errors(
    store: OciObjectStore,
    body: dict[str, Any],
    tool_name: str,
    artifact_type: str,
) -> tuple[list[str], dict[str, Any]]:
    errors: list[str] = []
    details: dict[str, Any] = {"type": artifact_type}
    key = _artifact_key(body, tool_name, artifact_type)
    details["key"] = key
    if not key:
        return [f"{tool_name} did not expose an artifact key"], details
    try:
        content = store.get(key)
    except Exception as exc:
        return [f"object storage reload failed for {key}: {exc}"], details
    details["bytes"] = len(content)
    if not content:
        errors.append(f"artifact {key} is empty")
        return errors, details

    try:
        if artifact_type == "pov":
            text = content.decode("utf-8")
            if "northwind" not in text.lower():
                errors.append("POV does not identify Northwind")
            if not key.endswith(".md"):
                errors.append(f"POV key is not Markdown: {key}")
        elif artifact_type == "poc_plan":
            payload = json.loads(content.decode("utf-8"))
            options = payload.get("poc_options") if isinstance(payload, dict) else None
            if not isinstance(options, list) or not options:
                errors.append("POC artifact has no options")
        elif artifact_type == "diagram":
            text = content.decode("utf-8")
            ElementTree.fromstring(text)
            lowered = text.lower()
            tier_hits = sum(
                marker in lowered
                for marker in ("iis", "web", "claims", "application", "oracle", "database")
            )
            if tier_hits < 3:
                errors.append("draw.io does not visibly represent the stated three tiers")
            if not key.endswith(".drawio"):
                errors.append(f"diagram key is not .drawio: {key}")
        elif artifact_type == "bom":
            workbook = openpyxl.load_workbook(BytesIO(content), data_only=False)
            cell_text = " ".join(
                str(cell.value or "")
                for sheet in workbook.worksheets
                for row in sheet.iter_rows()
                for cell in row
            ).lower()
            if "gen ai" in cell_text or "token" in cell_text:
                errors.append("BOM contains Gen AI/token line-item text")
            if not key.endswith(".xlsx"):
                errors.append(f"BOM key is not .xlsx: {key}")
        elif artifact_type == "jep":
            document = Document(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            lowered = text.lower()
            for marker in ("assessment", "build", "validate"):
                if marker not in lowered:
                    errors.append(f"JEP is missing {marker} phase")
            if not any(owner in lowered for owner in ("priya", "marcus", "jordan")):
                errors.append("JEP does not contain a grounded owner from the notes")
            if not any(target in lowered for target in ("200", "600", "45")):
                errors.append("JEP does not contain a grounded success criterion")
            if not key.endswith(".docx"):
                errors.append(f"JEP key is not .docx: {key}")
    except Exception as exc:
        errors.append(f"artifact {key} is not reloadable as {artifact_type}: {exc}")
    return errors, details


def _normalize_units(text: str) -> str:
    normalized = str(text or "").lower()
    replacements = (
        (r"\bmilliseconds?\b", "ms"),
        (r"\bconcurrent\s+users?\b", "users"),
        (r"\brequests?\s+per\s+second\b", "rps"),
        (r"\bgigabytes?\b", "gb"),
        (r"\bterabytes?\b", "tb"),
    )
    for pattern, replacement in replacements:
        normalized = re.sub(pattern, replacement, normalized)
    normalized = re.sub(r"(?<=\d)\s+(?=(?:ms|gb|tb|users|rps|%)(?:\b|$))", "", normalized)
    return re.sub(r"\s+", " ", normalized).strip()


def _claim_clause(text: str, start: int, end: int) -> str:
    left = max(text.rfind(mark, 0, start) for mark in (".", "!", "?", "\n", ";"))
    rights = [position for mark in (".", "!", "?", "\n", ";") if (position := text.find(mark, end)) >= 0]
    right = min(rights) if rights else len(text)
    return text[left + 1 : right].lower()


def _is_hedged_advisory(clause: str) -> bool:
    hedged = re.search(r"\b(?:typically|often|roughly|about|approximately|up to|in general)\b|~", clause)
    engagement_claim = re.search(
        r"\b(?:northwind|this (?:deal|engagement|customer)|for (?:you|them|northwind)|"
        r"their (?:measured|actual|current)|will (?:save|cost|achieve|deliver))\b",
        clause,
    )
    return bool(hedged and not engagement_claim)


def _measurement_is_grounded(value: str, grounded_text: str) -> bool:
    normalized_value = _normalize_units(value).replace(",", "")
    normalized_ground = _normalize_units(grounded_text).replace(",", "")
    if normalized_value in normalized_ground:
        return True
    number = re.search(r"\d+(?:\.\d+)?", normalized_value)
    if not number:
        return False
    claim_number = float(number.group(0))
    unit_match = re.search(r"\$|%|\b(?:ms|minutes|hours|days|ocpu|gb|tb|users|rps)\b", normalized_value)
    unit = unit_match.group(0) if unit_match else ""
    candidate_pattern = (
        rf"(?:\$\s*)?(\d+(?:\.\d+)?)" if unit == "$" else
        rf"(\d+(?:\.\d+)?)\s*{re.escape(unit)}\b" if unit else
        r"(\d+(?:\.\d+)?)"
    )
    for candidate in re.finditer(candidate_pattern, normalized_ground):
        grounded_number = float(candidate.group(1))
        tolerance = max(0.02 * abs(grounded_number), 0.5)
        if abs(claim_number - grounded_number) <= tolerance:
            return True
    return False


def _fabrication_errors(
    reply: str, grounded_text: str, artifact_text: str = ""
) -> list[str]:
    errors: list[str] = []
    reply_lower = reply.lower()
    grounded_lower = f"{grounded_text}\n{artifact_text}".lower()
    if re.search(r"\b(?:e5|e6)(?:\.flex)?\b[^.!?\n]{0,50}\b(?:ampere|arm)\b", reply_lower):
        errors.append("reply incorrectly describes E5/E6 as Ampere/Arm")

    precise_patterns = (
        r"\$\s*\d[\d,.]*",
        r"\b\d+(?:\.\d+)?\s*%",
        r"\b\d+(?:\.\d+)?\s*(?:ms|milliseconds|minutes|hours|days|ocpu|gb|tb)\b",
        r"\b\d+(?:\.\d+)?\s*(?:concurrent users|requests per second|rps)\b",
    )
    for pattern in precise_patterns:
        for match in re.finditer(pattern, reply_lower, re.IGNORECASE):
            value = re.sub(r"\s+", " ", match.group(0)).strip()
            clause = _claim_clause(reply_lower, match.start(), match.end())
            if _is_hedged_advisory(clause):
                continue
            if not _measurement_is_grounded(value, grounded_lower):
                errors.append(f"unsupported precise claim: {match.group(0)}")

    for match in re.finditer(
        r"\b([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,3})\s+"
        r"(?:achieved|reported|reduced|saved|improved|demonstrated)",
        reply,
    ):
        if match.group(1).lower() not in grounded_lower:
            errors.append(f"unsupported customer evidence: {match.group(1)}")
    return list(dict.fromkeys(errors))


def _artifact_grounding_text(
    store: OciObjectStore,
    body: dict[str, Any],
    tool_name: str,
    artifact_type: str,
) -> str:
    key = _artifact_key(body, tool_name, artifact_type)
    if not key:
        return ""
    try:
        content = store.get(key)
        if artifact_type == "bom":
            workbook = openpyxl.load_workbook(BytesIO(content), data_only=True)
            return " ".join(
                str(cell.value)
                for sheet in workbook.worksheets
                for row in sheet.iter_rows()
                for cell in row
                if cell.value not in (None, "")
            )
        if artifact_type == "jep":
            document = Document(BytesIO(content))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        return content.decode("utf-8", errors="replace")
    except Exception:
        return ""


def _needs_input_errors(turn: dict[str, Any], call: dict[str, Any]) -> list[str]:
    if str(call.get("result_status") or "") != "needs_input":
        return []
    if turn.get("required_inputs_supplied", True):
        return ["producer returned needs_input despite supplied grounding"]
    return []


def _artifact_prose_errors(reply: str) -> list[str]:
    errors: list[str] = []
    if re.search(r"\|\s*SKU\b", reply, re.IGNORECASE):
        errors.append("lookup reply contains a prose BOM table")
    if re.search(r"\$\s*\d[\d,.]*\s*(?:/\s*mo|per month|monthly)", reply, re.IGNORECASE):
        errors.append("lookup reply contains an invented monthly price")
    if re.search(r"\bfake\s+v\d+|\bversion\s+v\d+", reply, re.IGNORECASE):
        errors.append("lookup reply contains an ungrounded version string")
    return errors


def _behavior_errors(
    turn: dict[str, Any],
    body: dict[str, Any],
    prior_artifact_types: set[str],
) -> list[str]:
    errors: list[str] = []
    tools = _tool_names(body)
    tool_set = set(tools)
    generated = {tool for tool in tool_set if tool.startswith("generate_")}
    reply = str(body.get("reply") or "")
    kind = turn["kind"]

    if kind in {"conversation", "logistics"}:
        if generated:
            errors.append(f"conversational turn fired generate tools: {sorted(generated)}")
    elif kind.startswith("lookup"):
        if not (tool_set & LOOKUP_TOOLS):
            errors.append(f"lookup turn did not call a fetch tool: {tools}")
        if generated:
            errors.append(f"lookup turn fired generate tools: {sorted(generated)}")
        errors.extend(_artifact_prose_errors(reply))
        if kind == "lookup_empty" and not re.search(
            r"\b(?:no|none|not yet|haven't|have not|don't have|do not have)\b",
            reply,
            re.IGNORECASE,
        ):
            errors.append("empty lookup did not clearly say no artifact exists")
        if kind == "lookup_bom":
            if "token" not in reply.lower() and "gen ai" not in reply.lower():
                errors.append("BOM lookup did not answer the Gen AI/token question")
        if kind == "lookup_artifacts":
            for artifact_type in sorted(prior_artifact_types & {"pov", "diagram", "bom", "jep"}):
                if artifact_type not in reply.lower() and not (
                    artifact_type == "diagram" and "draw" in reply.lower()
                ):
                    errors.append(f"final lookup omitted real {artifact_type} artifact")
    elif kind == "deliverable":
        expected = str(turn["tool"])
        if expected not in tool_set:
            errors.append(f"expected {expected}, got {tools}")
        unexpected = generated - {expected}
        if unexpected:
            errors.append(f"deliverable turn fired unrelated generators: {sorted(unexpected)}")
        errors.extend(_needs_input_errors(turn, _call_for(body, expected)))
    elif kind == "selection":
        if "generate_poc_plan" not in tool_set:
            errors.append(f"POC selection was not recorded through generate_poc_plan: {tools}")
        unexpected = generated - {"generate_poc_plan"}
        if unexpected:
            errors.append(f"POC selection fired downstream generators: {sorted(unexpected)}")
        call = _call_for(body, "generate_poc_plan")
        data = call.get("result_data") if isinstance(call.get("result_data"), dict) else {}
        selected = data.get("selected_option") if isinstance(data, dict) else None
        if not isinstance(selected, dict) or not selected:
            errors.append("POC selection result did not record a selected option")
    elif kind == "notes":
        if generated:
            errors.append(f"notes turn fired generate tools: {sorted(generated)}")
        if not re.search(r"\b(?:note|notes|saved|captured|extracted|northwind)\b", reply, re.IGNORECASE):
            errors.append("notes turn did not acknowledge or summarize the uploaded notes")
    elif kind == "memory":
        if generated:
            errors.append(f"memory recall fired generate tools: {sorted(generated)}")
        required_groups = (
            ("northwind",),
            (".net", "member portal"),
            ("iis", "web"),
            ("claims",),
            ("oracle",),
            ("hipaa",),
            ("chicago", "us-chicago-1"),
        )
        for group in required_groups:
            if not any(item in reply.lower() for item in group):
                errors.append(f"memory reply omitted earlier fact: {'/'.join(group)}")
    return errors


def run(args: argparse.Namespace) -> dict[str, Any]:
    store = _store_from_config()
    run_id = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    customer_id = args.customer_id or f"native-sim-northwind-{run_id.lower()}"
    evidence: dict[str, Any] = {
        "scenario": "Northwind Health native engagement simulation",
        "run_id": run_id,
        "customer_id": customer_id,
        "base_url": args.base_url,
        "agent_mode": "native",
        "model": "Grok 4 via configured OCI Generative AI endpoint",
        "started_at": datetime.now(timezone.utc).isoformat(),
        "notes_text": NOTES,
        "turns": [],
        "artifact_reloads": [],
        "failures": [],
    }
    grounded_parts: list[str] = []
    artifact_grounding_parts: list[str] = []
    artifact_types: set[str] = set()
    max_latency = 0.0
    prior_chat_started = 0.0

    for turn in TURNS:
        turn_failures: list[str] = []
        upload_evidence: dict[str, Any] | None = None
        if turn.get("upload_notes"):
            try:
                upload, upload_elapsed, upload_status = _upload_note(
                    args.base_url,
                    customer_id,
                    args.timeout,
                )
                upload_evidence = {
                    "http_status": upload_status,
                    "elapsed_seconds": round(upload_elapsed, 3),
                    "response": upload,
                }
                max_latency = max(max_latency, upload_elapsed)
                if upload_status >= 500:
                    turn_failures.append(f"notes upload returned HTTP {upload_status}")
                if upload.get("extraction_status") != "ok":
                    turn_failures.append(
                        f"notes extraction status was {upload.get('extraction_status')!r}"
                    )
                grounded_parts.append(NOTES)
            except Exception as exc:
                upload_evidence = {"error": str(exc)}
                turn_failures.append(f"notes upload failed: {exc}")

        payload = {
            "customer_id": customer_id,
            "customer_name": "Northwind Health",
            "message": turn["message"],
            "session_id": turn["session_id"],
            "project_id": "native-engagement-sim",
            "project_name": "Native Engagement Simulation",
        }
        if prior_chat_started:
            wait_seconds = args.min_turn_interval - (time.monotonic() - prior_chat_started)
            if wait_seconds > 0:
                time.sleep(wait_seconds)
        prior_chat_started = time.monotonic()
        body: dict[str, Any] = {}
        status = 0
        elapsed = 0.0
        try:
            body, elapsed, status = _post_json(
                args.base_url,
                "/api/chat",
                payload,
                args.timeout,
            )
            max_latency = max(max_latency, elapsed)
            if status >= 500:
                turn_failures.append(f"chat returned HTTP {status}: {body}")
            elif status != 200:
                turn_failures.append(f"chat returned HTTP {status}: {body}")
        except Exception as exc:
            turn_failures.append(f"chat request failed: {exc}")

        grounded_parts.append(str(turn["message"]))
        if status == 200:
            turn_failures.extend(_behavior_errors(turn, body, artifact_types))
            if turn["id"] in {1, 2, 5, 7}:
                turn_failures.extend(
                    _fabrication_errors(
                        str(body.get("reply") or ""),
                        "\n".join(grounded_parts),
                        "\n".join(artifact_grounding_parts),
                    )
                )

            if turn["kind"] == "deliverable":
                expected_tool = str(turn["tool"])
                artifact_type = str(turn["artifact_type"])
                expected_call = _call_for(body, expected_tool)
                accepted_needs_input = (
                    str(expected_call.get("result_status") or "") == "needs_input"
                    and not turn.get("required_inputs_supplied", True)
                )
                artifact_errors: list[str] = []
                artifact_details: dict[str, Any] = {"type": artifact_type}
                if not accepted_needs_input:
                    artifact_errors, artifact_details = _artifact_validation_errors(
                        store,
                        body,
                        expected_tool,
                        artifact_type,
                    )
                download_ok = True
                download_detail = "not exposed through artifact_manifest"
                if not accepted_needs_input and artifact_type in {"diagram", "bom", "jep"}:
                    download_ok, download_detail = _download_manifest_artifact(
                        args.base_url,
                        body,
                        artifact_type,
                        args.timeout,
                    )
                    if not download_ok:
                        artifact_errors.append(download_detail)
                artifact_details["http_download"] = {
                    "ok": download_ok,
                    "detail": download_detail,
                }
                artifact_details["pass"] = not artifact_errors
                artifact_details["failures"] = artifact_errors
                evidence["artifact_reloads"].append(artifact_details)
                turn_failures.extend(artifact_errors)
                if not artifact_errors and not accepted_needs_input:
                    artifact_types.add(artifact_type)
                    artifact_text = _artifact_grounding_text(
                        store, body, expected_tool, artifact_type
                    )
                    if artifact_text:
                        artifact_grounding_parts.append(artifact_text)

        turn_record = {
            "turn": turn["id"],
            "session_id": turn["session_id"],
            "intent": turn["kind"],
            "message": turn["message"],
            "http_status": status,
            "elapsed_seconds": round(elapsed, 3),
            "reply": body.get("reply", "") if isinstance(body, dict) else "",
            "tool_calls": _compact_tool_calls(body),
            "artifact_manifest": body.get("artifact_manifest", {}) if isinstance(body, dict) else {},
            "upload": upload_evidence,
            "verdict": "PASS" if not turn_failures else "FAIL",
            "failures": list(dict.fromkeys(turn_failures)),
        }
        evidence["turns"].append(turn_record)
        for failure in turn_record["failures"]:
            evidence["failures"].append({"turn": turn["id"], "cause": failure})
        print(
            f"Turn {turn['id']:02d} [{turn_record['verdict']}] "
            f"{elapsed:.3f}s tools={_tool_names(body)}"
        )
        for failure in turn_record["failures"]:
            print(f"  - {failure}")

    evidence["max_turn_latency_seconds"] = round(max_latency, 3)
    evidence["completed_at"] = datetime.now(timezone.utc).isoformat()
    evidence["overall_verdict"] = "PASS" if not evidence["failures"] else "FAIL"
    args.report.parent.mkdir(parents=True, exist_ok=True)
    args.report.write_text(json.dumps(evidence, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(
        f"Overall: {evidence['overall_verdict']} | "
        f"max latency: {evidence['max_turn_latency_seconds']:.3f}s | report: {args.report}"
    )
    return evidence


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--base-url", default="http://127.0.0.1:18080")
    parser.add_argument("--report", type=Path, default=DEFAULT_REPORT)
    parser.add_argument("--customer-id", default="")
    parser.add_argument("--timeout", type=float, default=300.0)
    parser.add_argument(
        "--min-turn-interval",
        type=float,
        default=75.0,
        help="Minimum seconds between chat requests to stay below live model TPM limits.",
    )
    return parser.parse_args()


if __name__ == "__main__":
    result = run(_parse_args())
    raise SystemExit(0 if result["overall_verdict"] == "PASS" else 1)
