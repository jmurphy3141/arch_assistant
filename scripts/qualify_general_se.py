#!/usr/bin/env python3
"""Run reproducible multi-turn General SE qualification against an isolated stack."""
from __future__ import annotations

import argparse
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from io import BytesIO
import json
from pathlib import Path
import re
import time
from typing import Any
from urllib import request
import uuid
import sys

import yaml
import openpyxl
from docx import Document
from xml.etree import ElementTree

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from agent import consistency_contract, context_store
from agent.tools.bom import _extract_explicit_bom_inputs
from agent.object_store_oci import OciObjectStore


DEFAULT_REPORT = ROOT / "docs" / "general-se-live-qualification.json"
COMPLEX_REPORT = ROOT / "docs" / "general-se-complex-live-qualification.json"


@dataclass(frozen=True)
class Profile:
    name: str
    workload: str
    region: str
    services: tuple[str, ...]
    sizing: str
    duration: int
    criteria: tuple[str, ...]
    owners: str
    correction: str
    forbidden: tuple[str, ...] = ()
    upload_note: bool = False
    selection: int = 1
    tier_layout: str = ""


PROFILES = (
    Profile(
        "Granite Harbor Credit Union",
        "an on-premises Java payments application backed by PostgreSQL",
        "us-ashburn-1",
        ("VM.Standard.E5.Flex", "PostgreSQL", "Object Storage", "Logging", "Site-to-Site VPN"),
        "2 E5 application servers, 2 OCPU and 16 GB memory each; PostgreSQL 2 OCPU; 1 TB Object Storage",
        15,
        ("p95 under 500 milliseconds at 200 requests per second", "restore within 60 minutes"),
        "Oracle SA and Granite Harbor technical lead each commit 10 hours per week",
        "Correction: connectivity is Site-to-Site VPN, not FastConnect.",
        ("OCI WAF", "Flexible Load Balancer"),
    ),
    Profile(
        "BrightPath Clinics",
        "a patient scheduling API on private virtual machines backed by PostgreSQL",
        "us-chicago-1",
        ("Flexible Load Balancer", "VM.Standard.E5.Flex", "PostgreSQL", "Logging", "Monitoring"),
        "2 E5 API servers, 2 OCPU and 16 GB memory each; PostgreSQL 2 OCPU",
        12,
        ("p95 below 400 milliseconds", "99.9% availability"),
        "Oracle SA and BrightPath platform lead each commit 10 hours per week",
        "Correction: the load balancer is public, while the API and database remain private.",
        ("OCI WAF",),
        True,
    ),
    Profile(
        "VoltForge Manufacturing",
        "a VMware-hosted manufacturing telemetry service with no database",
        "us-phoenix-1",
        ("Object Storage", "Block Volume", "Site-to-Site VPN", "Logging", "Monitoring"),
        "2 application servers, 2 OCPU and 16 GB memory each; 500 GB Block Volume; 1 TB Object Storage",
        15,
        ("process 300 requests per second", "recover within 30 minutes"),
        "Oracle SA and VoltForge engineer each commit 8 hours per week",
        "Correction: compute family was not selected; use the standard E6 Flex default for the POC.",
        ("PostgreSQL", "OCI WAF", "Flexible Load Balancer"),
    ),
    Profile(
        "AeroSpan Freight",
        "a public order-routing service backed by PostgreSQL",
        "us-ashburn-1",
        ("Flexible Load Balancer", "VM.Standard.E5.Flex", "PostgreSQL", "Object Storage", "Monitoring"),
        "2 E5 routing servers, 2 OCPU and 16 GB memory each; PostgreSQL 2 OCPU; 1 TB Object Storage",
        15,
        ("sustain 200 requests per second", "p95 under 350 milliseconds"),
        "Oracle SA and AeroSpan technical lead each commit 10 hours per week",
        "Correction: this is internet-facing through the Flexible Load Balancer, but WAF is out of scope.",
        ("OCI WAF",),
    ),
    Profile(
        "Sterling Mutual",
        "a private claims batch-processing application backed by PostgreSQL",
        "us-chicago-1",
        ("VM.Standard.E5.Flex", "PostgreSQL", "Block Volume", "Logging"),
        "2 E5 batch servers, 2 OCPU and 16 GB memory each; PostgreSQL 2 OCPU; 500 GB Block Volume",
        10,
        ("finish the batch within 2 hours", "restore within 45 minutes"),
        "Oracle SA and Sterling Mutual owner each commit 9 hours per week",
        "Correction: the workload is private-only with no public edge.",
        ("Object Storage", "OCI WAF", "Flexible Load Balancer"),
    ),
)


COMPLEX_PROFILES = (
    Profile(
        "Northstar Outfitters",
        "a revenue-critical Java ecommerce platform with separate NGINX web, Java application, and PostgreSQL database tiers",
        "us-ashburn-1",
        ("OCI WAF", "Flexible Load Balancer", "VM.Standard.E5.Flex", "PostgreSQL DB System", "Object Storage", "FastConnect", "Logging", "Monitoring"),
        "2 E5 web servers at 2 OCPU and 16 GB each; 4 E5 application servers at 4 OCPU and 32 GB each; HA PostgreSQL DB System at 4 OCPU; 2 TB Object Storage",
        20,
        ("sustain 750 orders per minute", "keep p95 checkout latency under 450 milliseconds", "recover within 30 minutes"),
        "Oracle SA, Northstar application lead, and Northstar network engineer each commit 12 hours per week",
        "Correction: production connectivity is redundant FastConnect, not Site-to-Site VPN. Treat FastConnect as authoritative.",
        ("Site-to-Site VPN", "Container Engine for Kubernetes (OKE)"),
        True,
        2,
        "Internet traffic enters OCI WAF and a public Flexible Load Balancer; the web tier is private, the application tier is isolated in a separate private subnet, and PostgreSQL is in a database-only private subnet. On-premises ERP traffic uses redundant FastConnect through a DRG.",
    ),
    Profile(
        "Meridian Health Plans",
        "a HIPAA-regulated .NET member portal with distinct IIS web, claims-service application, and Oracle database tiers",
        "us-chicago-1",
        ("OCI WAF", "Flexible Load Balancer", "VM.Standard.E5.Flex", "Oracle Base Database Service", "File Storage", "FastConnect", "OCI IAM", "Audit Logging", "Monitoring"),
        "2 E5 IIS web servers at 2 OCPU and 16 GB each; 3 E5 claims application servers at 4 OCPU and 32 GB each; HA Oracle Base Database Service at 4 OCPU; 1 TB File Storage",
        25,
        ("support 400 concurrent portal sessions", "keep p95 claims lookup under 600 milliseconds", "restore service within 45 minutes"),
        "Oracle SA, Meridian security lead, and Meridian application owner each commit 10 hours per week",
        "Correction: the database target is Oracle Base Database Service, not PostgreSQL. Preserve the three private tiers and FastConnect.",
        ("PostgreSQL DB System", "Site-to-Site VPN", "Object Storage"),
        False,
        3,
        "OCI WAF fronts a public Flexible Load Balancer; IIS web, claims application, and Oracle database run in separate private subnets. Private hospital and identity-system integration traverses FastConnect and a DRG; audit evidence is mandatory.",
    ),
    Profile(
        "Ironvale Components",
        "a private supplier-order portal with separate Apache web, Spring Boot application, and PostgreSQL database tiers",
        "us-phoenix-1",
        ("Flexible Load Balancer", "VM.Standard.E6.Flex", "PostgreSQL DB System", "Block Volume", "FastConnect", "Logging", "Monitoring"),
        "2 E6 Apache web servers at 2 OCPU and 16 GB each; 2 E6 Spring Boot application servers at 4 OCPU and 32 GB each; HA PostgreSQL DB System at 4 OCPU; 1 TB Block Volume",
        18,
        ("process 250 supplier orders per minute", "keep p95 API response under 500 milliseconds", "recover within 60 minutes"),
        "Oracle SA, Ironvale platform owner, and Ironvale network lead each commit 9 hours per week",
        "Correction: partner access remains private through FastConnect; do not add an internet edge or WAF.",
        ("OCI WAF", "Site-to-Site VPN", "Object Storage"),
        False,
        1,
        "A private Flexible Load Balancer receives partner traffic through FastConnect and a DRG. Apache web, Spring Boot application, and PostgreSQL database tiers occupy three separate private subnets with no public ingress.",
    ),
    Profile(
        "BlueCurrent Resorts",
        "a public hotel-booking platform with separate web storefront, reservation-service application, and PostgreSQL database tiers",
        "eu-frankfurt-1",
        ("OCI WAF", "Flexible Load Balancer", "VM.Standard.E5.Flex", "PostgreSQL DB System", "Object Storage", "FastConnect", "Logging", "Monitoring"),
        "2 E5 storefront web servers at 2 OCPU and 16 GB each; 4 E5 reservation application servers at 4 OCPU and 32 GB each; HA PostgreSQL DB System at 4 OCPU; 3 TB Object Storage",
        21,
        ("sustain 500 booking requests per minute", "keep p95 search latency under 400 milliseconds", "recover within 30 minutes"),
        "Oracle SA, BlueCurrent digital owner, and BlueCurrent database lead each commit 11 hours per week",
        "Correction: hotel property-system integration uses FastConnect; only guest booking traffic is public through WAF.",
        ("Site-to-Site VPN", "Oracle Base Database Service"),
        True,
        2,
        "Guest traffic enters OCI WAF and a public Flexible Load Balancer. Web storefront, reservation application, and PostgreSQL database tiers are isolated in separate private subnets; property systems connect through FastConnect and a DRG.",
    ),
    Profile(
        "Cedar Ridge University",
        "a student-services portal with separate web presentation, enrollment application, and Oracle database tiers",
        "us-ashburn-1",
        ("OCI WAF", "Flexible Load Balancer", "VM.Standard.E6.Flex", "Oracle Base Database Service", "File Storage", "FastConnect", "OCI IAM", "Logging", "Monitoring"),
        "2 E6 presentation web servers at 2 OCPU and 16 GB each; 3 E6 enrollment application servers at 4 OCPU and 32 GB each; HA Oracle Base Database Service at 4 OCPU; 2 TB File Storage",
        24,
        ("support 1200 concurrent enrollment users", "keep p95 registration under 700 milliseconds", "restore within 45 minutes"),
        "Oracle SA, Cedar Ridge applications director, and campus network engineer each commit 10 hours per week",
        "Correction: the web tier is public only through OCI WAF and the load balancer; application and database tiers have no public IPs.",
        ("Site-to-Site VPN", "PostgreSQL DB System", "Object Storage"),
        False,
        3,
        "OCI WAF and a public Flexible Load Balancer serve the web tier. The enrollment application and Oracle database occupy separate private subnets, while campus identity and records systems use redundant FastConnect through a DRG.",
    ),
)


def _post_json(base_url: str, path: str, payload: dict[str, Any]) -> tuple[dict[str, Any], float]:
    started = time.monotonic()
    req = request.Request(
        base_url + path,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with request.urlopen(req, timeout=65) as response:
        body = json.loads(response.read().decode("utf-8"))
    return body, time.monotonic() - started


def _upload_note(base_url: str, customer_id: str, text: str) -> dict[str, Any]:
    boundary = "----archie-qualification-" + uuid.uuid4().hex
    fields = [
        ("customer_id", customer_id),
        ("note_name", "discovery-note.txt"),
    ]
    chunks: list[bytes] = []
    for name, value in fields:
        chunks.append(
            f"--{boundary}\r\nContent-Disposition: form-data; name=\"{name}\"\r\n\r\n{value}\r\n".encode()
        )
    chunks.append(
        f"--{boundary}\r\nContent-Disposition: form-data; name=\"file\"; filename=\"discovery-note.txt\"\r\n"
        "Content-Type: text/plain\r\n\r\n".encode()
        + text.encode()
        + b"\r\n"
    )
    chunks.append(f"--{boundary}--\r\n".encode())
    req = request.Request(
        base_url + "/api/notes/upload",
        data=b"".join(chunks),
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    with request.urlopen(req, timeout=65) as response:
        return json.loads(response.read().decode())


def _tool_names(body: dict[str, Any]) -> list[str]:
    return [str(call.get("tool") or "") for call in body.get("tool_calls", []) or []]


def _call_for(body: dict[str, Any], tool: str) -> dict[str, Any]:
    return next((call for call in body.get("tool_calls", []) or [] if call.get("tool") == tool), {})


def _store_from_config() -> OciObjectStore:
    cfg = yaml.safe_load((ROOT / "config.yaml").read_text())
    persistence = cfg["persistence"]
    return OciObjectStore(
        region=persistence["region"],
        namespace=persistence["namespace"],
        bucket_name=persistence["bucket_name"],
    )


def _turns(profile: Profile) -> list[tuple[str, str]]:
    services = ", ".join(profile.services)
    criteria = ", ".join(profile.criteria)
    return [
        ("introduction", f"We're beginning discovery with {profile.name}. They run {profile.workload}, but I do not have the full design yet."),
        ("advice", "What architecture direction would you consider first, and what tradeoffs should we discuss? Do not generate an artifact."),
        ("note", f"Please save this discovery note only: primary workload is {profile.workload}; target services currently include {services}. {profile.tier_layout} Do not generate artifacts."),
        ("discovery", f"The approved OCI region is {profile.region}. The POC duration is {profile.duration} days. {profile.owners}."),
        ("criteria", f"In-scope services are {services}. Required topology: {profile.tier_layout or 'use the confirmed workload topology'}. Success criteria: {criteria}. Out of scope: {', '.join(profile.forbidden) or 'production cutover'}."),
        ("correction", profile.correction + " Treat this as the current authoritative fact."),
        ("recall", "Recall the current workload, region, services, exclusions, owners, duration, and success criteria. Do not generate anything."),
        ("poc", "Create three POC options now using the accumulated engagement context; do not ask me to restate the fact sheet."),
        ("selection", f"Option {profile.selection}"),
        ("bom", f"Build only the BOM/XLSX for the selected POC. Sizing: {profile.sizing}. Use {profile.region}. Do not generate a diagram or JEP."),
        ("bom_discussion", "Explain one important sizing or cost decision in that BOM. Do not regenerate or create another artifact."),
        ("diagram", "Generate only the draw.io Diagram from the selected POC and finalized BOM. Do not generate Terraform, WAF, or JEP."),
        ("jep", f"Generate only the JEP artifact for the selected POC, finalized BOM, and Diagram. Use {profile.duration} days: Phase 1 Assessment days 1-3, Phase 2 Build days 4-{profile.duration - 4}, Phase 3 Validate days {profile.duration - 3}-{profile.duration}. Preserve the authoritative owners and success criteria; include at least three risks. The named Oracle SA and customer application owner jointly approve go/no-go. If no-go, stop testing, export the evidence, and remove POC resources without production cutover."),
    ]


_COMPUTE_OCPU_SKUS = {"B93113", "B97384", "B111129", "B94176", "B93297"}
_COMPUTE_MEMORY_SKUS = {"B93114", "B97385", "B111130", "B94177", "B93298"}


def _service_present(service: str, text: str) -> bool:
    service_id = consistency_contract.canonical_service_id(service)
    aliases = next((aliases for candidate, aliases in consistency_contract._SERVICE_ALIASES if candidate == service_id), ())
    return any(alias in text.lower() for alias in aliases) if aliases else service.lower() in text.lower()


def _xlsx_rows(content: bytes) -> tuple[list[dict[str, Any]], str]:
    workbook = openpyxl.load_workbook(BytesIO(content), data_only=False)
    sheet = workbook["BOM"]
    headers = [str(cell.value or "").strip() for cell in sheet[1]]
    rows = [
        dict(zip(headers, values))
        for values in sheet.iter_rows(min_row=2, values_only=True)
        if any(value not in (None, "") for value in values)
        and str(values[0] or "").strip().upper() != "TOTAL"
    ]
    scope_text = ""
    if "Scope Coverage" in workbook.sheetnames:
        scope_text = " ".join(
            str(value or "")
            for row in workbook["Scope Coverage"].iter_rows(values_only=True)
            for value in row
        )
    return rows, scope_text


def _assert_bom_content(content: bytes, profile: Profile) -> list[str]:
    errors: list[str] = []
    rows, scope_text = _xlsx_rows(content)
    sizing = _extract_explicit_bom_inputs(profile.sizing)
    expected_ocpu = float((sizing.get("compute") or {}).get("ocpu") or 0)
    expected_memory = float((sizing.get("memory") or {}).get("gb") or 0)
    produced_ocpu = sum(float(row.get("Quantity") or 0) for row in rows if str(row.get("SKU") or "").upper() in _COMPUTE_OCPU_SKUS)
    produced_memory = sum(float(row.get("Quantity") or 0) for row in rows if str(row.get("SKU") or "").upper() in _COMPUTE_MEMORY_SKUS)
    expected_instances = int((sizing.get("compute") or {}).get("instance_count") or 0)
    produced_instances = sum(int(row.get("Instance Count") or 1) for row in rows if str(row.get("SKU") or "").upper() in _COMPUTE_OCPU_SKUS)
    if abs(produced_ocpu - expected_ocpu) > 0.0001:
        errors.append(f"compute OCPU expected {expected_ocpu:g}, got {produced_ocpu:g}")
    if abs(produced_memory - expected_memory) > 0.0001:
        errors.append(f"compute memory expected {expected_memory:g} GB, got {produced_memory:g} GB")
    if expected_instances and produced_instances != expected_instances:
        errors.append(f"compute instances expected {expected_instances}, got {produced_instances}")
    workbook_text = " ".join(
        " ".join(str(row.get(key) or "") for key in ("Description", "Notes", "SKU"))
        for row in rows
    ) + " " + scope_text
    represented = {
        service_id
        for service_id, aliases in consistency_contract._SERVICE_ALIASES
        if any(alias in workbook_text.lower() for alias in aliases)
    }
    for service in profile.services:
        service_id = consistency_contract.canonical_service_id(service)
        if service_id and service_id not in represented:
            errors.append(f"required service missing: {service}")
    database = sizing.get("database") or {}
    if database:
        database_id = str(database.get("service_id") or "")
        sized_rows = [
            row
            for row in rows
            if consistency_contract.canonical_service_id(f"{row.get('Description', '')} {row.get('Notes', '')}") == database_id
            and "ocpu" in str(row.get("Metric") or row.get("Description") or "").lower()
        ]
        sized = any(
            float(row.get("Quantity") or 0) > 0 for row in sized_rows
        )
        if not sized:
            errors.append(f"database is not sized: {database_id or database.get('display_name')}")
        else:
            produced_db_ocpu = sum(float(row.get("Quantity") or 0) for row in sized_rows)
            if abs(produced_db_ocpu - float(database.get("ocpu") or 0)) > 0.0001:
                errors.append(f"database OCPU expected {float(database.get('ocpu')):g}, got {produced_db_ocpu:g}")
    for field, sku, label in (
        ("block_gb", "B91961", "Block Volume"),
        ("object_gb", "B91628", "Object Storage"),
        ("file_gb", "BFILE01", "File Storage"),
    ):
        expected = (sizing.get("storage") or {}).get(field)
        if expected is None:
            continue
        produced = sum(float(row.get("Quantity") or 0) for row in rows if str(row.get("SKU") or "").upper() == sku)
        if abs(produced - float(expected)) > 0.0001:
            errors.append(f"{label} capacity expected {float(expected):g} GB, got {produced:g} GB")
    for forbidden in profile.forbidden:
        if _service_present(forbidden, workbook_text):
            errors.append(f"forbidden service present: {forbidden}")
    return errors


def _required_private_tier_count(tier_layout: str) -> int:
    text = str(tier_layout or "").lower()
    if not any(phrase in text for phrase in ("separate private subnet", "three private", "distinct private")):
        return 0
    return sum(bool(re.search(pattern, text)) for pattern in (r"\bweb\b|iis|apache|storefront|presentation", r"\bapp(?:lication)?\b|claims|spring boot|reservation|enrollment", r"\bdatabase\b|postgresql|oracle"))


def _drawio_labels(content: bytes) -> list[str]:
    root = ElementTree.fromstring(content)
    return [str(element.attrib.get("value") or "") for element in root.iter() if element.attrib.get("value")]


def _assert_diagram_content(content: bytes, profile: Profile) -> list[str]:
    labels = _drawio_labels(content)
    text = " ".join(labels).lower()
    private_subnets = sum(bool(re.search(r"\bprivate\b.*\bsubnet\b|\bsubnet\b.*\bprivate\b", label, re.IGNORECASE)) for label in labels)
    required_tiers = _required_private_tier_count(profile.tier_layout)
    errors = []
    if private_subnets < required_tiers:
        errors.append(f"private subnets expected at least {required_tiers}, got {private_subnets}")
    for service in profile.services:
        aliases = next((aliases for service_id, aliases in consistency_contract._SERVICE_ALIASES if service_id == consistency_contract.canonical_service_id(service)), ())
        if aliases and not any(alias in text for alias in aliases):
            errors.append(f"required service label missing: {service}")
    for forbidden in profile.forbidden:
        if _service_present(forbidden, text):
            errors.append(f"forbidden service label present: {forbidden}")
    return errors


def _jep_text(content: bytes, suffix: str) -> str:
    if suffix == ".md":
        return content.decode("utf-8", errors="replace")
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def _assert_jep_content(content: bytes, suffix: str, profile: Profile) -> list[str]:
    text = re.sub(r"\s+", " ", _jep_text(content, suffix).replace("–", "-").replace("—", "-")).lower()
    owner_clause = re.split(r"\beach commit\b", profile.owners, flags=re.IGNORECASE)[0]
    owners = [part.strip(" ,.") for part in re.split(r",|\band\b", owner_clause, flags=re.IGNORECASE) if part.strip(" ,.")]
    windows = ("days 1-3", f"days 4-{profile.duration - 4}", f"days {profile.duration - 3}-{profile.duration}")
    errors = [f"owner missing: {owner}" for owner in owners if owner.lower() not in text]
    errors.extend(f"phase window missing: {window}" for window in windows if window.lower() not in text)
    errors.extend(f"success criterion missing: {criterion}" for criterion in profile.criteria if criterion.lower() not in text)
    return errors


def _reply_grounding_errors(reply: str, engagement_context: str) -> list[str]:
    text = str(reply or "")
    grounded = str(engagement_context or "").lower()
    errors: list[str] = []
    if re.search(r"(?:e5|e6)(?:\.flex)?[^.!?\n]{0,60}\b(?:ampere|arm)\b|\b(?:ampere|arm)\b[^.!?\n]{0,60}(?:e5|e6)(?:\.flex)?", text, re.IGNORECASE):
        errors.append("E5/E6 incorrectly described as Ampere/Arm")
    for token in re.findall(r"\b\d+(?:\.\d+)?\s*%", text):
        if token.lower().replace(" ", "") not in grounded.replace(" ", ""):
            errors.append(f"unsupported percentage: {token}")
    for match in re.finditer(
        r"(?:[Cc]ustomer|[Cc]lient|[Cc]ase study|[Rr]eference)\s+(?:named\s+)?"
        r"([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,3})\s+"
        r"(?:achieved|reported|reduced|saved|improved|saw|demonstrated)",
        text,
    ):
        if match.group(1).lower() not in grounded:
            errors.append(f"unsupported customer evidence: {match.group(1)}")
    for match in re.finditer(
        r"\bsla\b[^.!?\n]{0,50}?\b(\d+(?:\.\d+)?%?)|"
        r"\b(\d+(?:\.\d+)?%?)\b[^.!?\n]{0,30}?\bsla\b",
        text,
        re.IGNORECASE,
    ):
        token = next((group for group in match.groups() if group), "")
        if token and token.lower().replace(" ", "") not in grounded.replace(" ", ""):
            errors.append(f"unsupported SLA figure: {token}")
    return list(dict.fromkeys(errors))


def run_profile(base_url: str, store: OciObjectStore, profile: Profile, run_id: str) -> dict[str, Any]:
    customer_id = f"general-se-{run_id}-{profile.name.lower().replace(' ', '-')[:28]}"
    evidence: dict[str, Any] = {
        "customer_id": customer_id,
        "profile": asdict(profile),
        "turns": [],
        "failures": [],
    }
    if profile.upload_note:
        upload = _upload_note(
            base_url,
            customer_id,
            f"Uploaded discovery: {profile.workload}. Region {profile.region}. Services: {', '.join(profile.services)}.",
        )
        evidence["upload"] = upload
        if upload.get("extraction_status") != "ok":
            evidence["failures"].append({"stage": "upload", "cause": upload})

    bodies: dict[str, dict[str, Any]] = {}
    engagement_grounding = json.dumps(asdict(profile), default=str) + " " + " ".join(message for _, message in _turns(profile))
    for stage, message in _turns(profile):
        try:
            body, elapsed = _post_json(base_url, "/api/chat", {
                "customer_id": customer_id,
                "customer_name": profile.name,
                "message": message,
                "project_id": run_id,
                "project_name": "General SE Live Qualification",
            })
        except Exception as exc:
            evidence["failures"].append({"stage": stage, "cause": f"request failed: {exc}"})
            break
        bodies[stage] = body
        tools = _tool_names(body)
        evidence["turns"].append({
            "stage": stage,
            "message": message,
            "elapsed_seconds": round(elapsed, 3),
            "reply": body.get("reply", ""),
            "tools": tools,
            "trace_id": body.get("trace_id", ""),
            "artifact_manifest": body.get("artifact_manifest", {}),
        })
        if elapsed > 60:
            evidence["failures"].append({"stage": stage, "cause": f"elapsed {elapsed:.3f}s exceeds 60s"})
        allowed = {
            "note": {"save_notes"}, "poc": {"generate_poc_plan"},
            "selection": {"generate_poc_plan"}, "bom": {"generate_bom"},
            "diagram": {"generate_diagram"}, "jep": {"generate_jep"},
        }.get(stage, set())
        if set(tools) - allowed:
            evidence["failures"].append({"stage": stage, "cause": f"unexpected tools: {tools}"})
        if stage in {"advice", "discovery", "criteria", "correction", "recall", "bom_discussion"}:
            question_count = str(body.get("reply") or "").count("?")
            if question_count > 3:
                evidence["failures"].append({"stage": stage, "cause": f"unbounded questions: {question_count}"})
        for error in _reply_grounding_errors(str(body.get("reply") or ""), engagement_grounding):
            evidence["failures"].append({"stage": stage, "cause": error})

    try:
        context = context_store.read_context(store, customer_id, profile.name)
        evidence["context_reload"] = {
            "engagement_summary": context_store.get_archie_state(context).get("engagement_summary", ""),
            "resolved_decisions": context_store.get_archie_state(context).get("resolved_decisions", {}),
        }
    except Exception as exc:
        evidence["failures"].append({"stage": "context_reload", "cause": str(exc)})
        context = {}

    poc_call = _call_for(bodies.get("poc", {}), "generate_poc_plan")
    options = (poc_call.get("result_data") or {}).get("poc_options", [])
    if len(options) != 3:
        evidence["failures"].append({"stage": "poc", "cause": f"expected 3 options, got {len(options)}"})
    selection_tools = _tool_names(bodies.get("selection", {}))
    if selection_tools != ["generate_poc_plan"]:
        evidence["failures"].append({"stage": "selection", "cause": f"selection tools: {selection_tools}"})

    selected = ((context_store.get_archie_state(context).get("resolved_decisions") or {}).get("poc") or {}).get("selected_option") or {}
    selected_text = " ".join(str(item) for item in selected.get("oci_services", []) or []).lower()
    for forbidden in profile.forbidden:
        if forbidden.lower() in selected_text:
            evidence["failures"].append({"stage": "poc", "cause": f"forbidden service selected: {forbidden}"})
    if profile.tier_layout:
        if "fastconnect" not in selected_text:
            evidence["failures"].append({"stage": "poc", "cause": "FastConnect missing from selected POC"})
        selected_name = str(selected.get("option_name") or "")
        expected_name = str((options[profile.selection - 1] if len(options) >= profile.selection else {}).get("option_name") or "")
        if selected_name != expected_name:
            evidence["failures"].append({"stage": "selection", "cause": f"expected option {profile.selection} ({expected_name}), got {selected_name}"})

    artifact_checks: list[dict[str, Any]] = []
    for stage, tool, suffixes in (
        ("bom", "generate_bom", (".xlsx",)),
        ("diagram", "generate_diagram", (".drawio",)),
        ("jep", "generate_jep", (".md", ".docx")),
    ):
        call = _call_for(bodies.get(stage, {}), tool)
        data = call.get("result_data") or {}
        keys = [call.get("artifact_key"), data.get("docx_key"), (data.get("bom_xlsx") or {}).get("key")]
        for suffix in suffixes:
            key = next((str(item) for item in keys if item and str(item).endswith(suffix)), "")
            content = store.get(key) if key and store.head(key) else b""
            ok = bool(content)
            content_errors: list[str] = []
            if content:
                try:
                    if suffix == ".xlsx":
                        content_errors = _assert_bom_content(content, profile)
                    elif suffix == ".drawio":
                        content_errors = _assert_diagram_content(content, profile)
                    else:
                        content_errors = _assert_jep_content(content, suffix, profile)
                except Exception as exc:
                    content_errors = [f"content parse failed: {exc}"]
            artifact_checks.append({
                "stage": stage,
                "suffix": suffix,
                "key": key,
                "direct_reload": ok,
                "content_assertions_enabled": True,
                "content_errors": content_errors,
                "content_valid": ok and not content_errors,
            })
            if not ok:
                evidence["failures"].append({"stage": stage, "cause": f"{suffix} failed direct reload: {key}"})
            for error in content_errors:
                evidence["failures"].append({"stage": stage, "cause": f"{suffix} {error}"})
    evidence["artifact_checks"] = artifact_checks
    evidence["passed"] = not evidence["failures"]
    return evidence


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--base-url", default="http://127.0.0.1:18080")
    parser.add_argument("--suite", choices=("general", "complex-three-tier"), default="general")
    parser.add_argument("--report", type=Path)
    args = parser.parse_args()
    profiles = COMPLEX_PROFILES if args.suite == "complex-three-tier" else PROFILES
    if args.report is None:
        args.report = COMPLEX_REPORT if args.suite == "complex-three-tier" else DEFAULT_REPORT
    run_id = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    store = _store_from_config()
    retained_failures: list[dict[str, Any]] = []
    if args.report.exists():
        try:
            previous = json.loads(args.report.read_text())
            retained_failures.extend(previous.get("retained_failures", []) or [])
            retained_failures.extend(
                engagement
                for engagement in previous.get("engagements", []) or []
                if not engagement.get("passed")
            )
        except (OSError, ValueError):
            pass
    report: dict[str, Any] = {
        "run_id": run_id,
        "base_url": args.base_url,
        "production_port_touched": False,
        "started_at": datetime.now(timezone.utc).isoformat(),
        "engagements": [],
        "retained_failures": retained_failures,
        "consecutive_passes": 0,
    }
    for profile in profiles:
        result = run_profile(args.base_url, store, profile, run_id)
        report["engagements"].append(result)
        report["consecutive_passes"] = report["consecutive_passes"] + 1 if result["passed"] else 0
        args.report.write_text(json.dumps(report, indent=2, default=str) + "\n")
        print(f"{profile.name}: {'PASS' if result['passed'] else 'FAIL'}; streak={report['consecutive_passes']}", flush=True)
        for failure in result["failures"]:
            print(f"  {failure['stage']}: {failure['cause']}", flush=True)
    report["completed_at"] = datetime.now(timezone.utc).isoformat()
    report["qualified"] = report["consecutive_passes"] >= 5
    args.report.write_text(json.dumps(report, indent=2, default=str) + "\n")
    return 0 if report["qualified"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
