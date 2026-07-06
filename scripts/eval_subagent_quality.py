#!/usr/bin/env python3
"""Measure OCI sub-agent artifact quality without changing producer behavior."""

from __future__ import annotations

import argparse
import asyncio
from collections import defaultdict
from datetime import datetime, timezone
from decimal import Decimal, InvalidOperation
from io import BytesIO
import inspect
import json
import math
from pathlib import Path
import re
import shutil
import statistics
import subprocess
import sys
import tempfile
import time
from typing import Any, Callable
import uuid
from xml.etree import ElementTree

import openpyxl
import yaml
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
RUBRIC_ROOT = ROOT / "eval" / "rubrics"
GOLDEN_ROOT = ROOT / "eval" / "golden"
DEFAULT_REPORT = ROOT / "docs" / "subagent-quality.json"
ARTIFACT_TYPES = ("pov", "jep", "bom", "diagram", "waf", "terraform")
_JUDGE_CONTENT_LIMIT = 30_000


FIXTURE: dict[str, Any] = {
    "customer_id": "quality-northwind",
    "customer_name": "Northwind Health",
    "engagement_id": "quality-northwind",
    "region": "us-chicago-1",
    "architecture_summary": (
        "An internet-facing three-tier .NET member claims portal in us-chicago-1: "
        "OCI WAF and a public Flexible Load Balancer front private E5 Flex web and "
        "application tiers; Oracle Base Database Service, Block Volume, Object "
        "Storage, Logging, Monitoring, IAM, Vault, NSGs, and private subnets protect "
        "HIPAA-regulated claims data."
    ),
    "facts": {
        "workload": "internet-facing three-tier .NET member claims portal",
        "current_platform": "on-premises",
        "region": "us-chicago-1",
        "compliance": ["HIPAA"],
        "competitive_alternative": "remain on-premises",
        "services": [
            "OCI Web Application Firewall",
            "OCI Flexible Load Balancer",
            "VM.Standard.E5.Flex",
            "Oracle Base Database Service",
            "Block Volume",
            "Object Storage",
            "Logging",
            "Monitoring",
            "IAM",
            "Vault",
        ],
        "sizing": {
            "web": "2 servers, 2 OCPU and 16 GB each",
            "application": "2 servers, 4 OCPU and 32 GB each",
            "database": "4 OCPU, HA",
            "block_storage_gb": 500,
            "object_storage_gb": 100,
        },
        "poc": {
            "duration": "6 weeks",
            "phases": [
                "Phase 1 Assessment — Week 1",
                "Phase 2 Build — Weeks 2-3",
                "Phase 3 Validate — Weeks 4-6",
            ],
            "success_criteria": [
                "p95 response time under 300 ms",
                "99.9% availability during the validation window",
                "recovery time under 60 minutes",
                "30% infrastructure cost reduction target within 12 months",
            ],
            "owners": [
                "Oracle solution architect — architecture and evidence",
                "Northwind application owner — test execution",
                "Northwind security lead — HIPAA control review",
            ],
        },
    },
    "structured_inputs": {
        "region": "us-chicago-1",
        "architecture_option": "OCI native three-tier claims portal",
        "strict_scope": True,
        "target_service_ids": [
            "security.waf",
            "network.load_balancer.flexible",
            "compute.vm.standard.e5.flex",
            "database.oracle",
            "storage.block",
            "storage.object",
            "observability.logging",
            "observability.monitoring",
        ],
        "target_services": [
            "OCI WAF",
            "Flexible Load Balancer",
            "VM.Standard.E5.Flex",
            "Oracle Base Database Service",
            "Block Volume",
            "Object Storage",
            "Logging",
            "Monitoring",
        ],
        "compute": {
            "ocpu": 12,
            "instance_count": 4,
            "tiers": [
                {"name": "web", "instance_count": 2, "ocpu": 4, "memory_gb": 32},
                {"name": "application", "instance_count": 2, "ocpu": 8, "memory_gb": 64},
            ],
        },
        "memory": {"gb": 96},
        "database": {
            "service_id": "database.oracle",
            "display_name": "Oracle Base Database Service",
            "ocpu": 4,
            "instance_count": 1,
            "ha": True,
        },
        "storage": {
            "block_tb": 0.48828125,
            "block_gb": 500,
            "object_tb": 0.09765625,
            "object_gb": 100,
        },
        "workloads": [".NET member claims portal"],
        "assumptions": ["730 hours per month", "estimate only and non-binding"],
    },
}


TASKS = {
    "pov": (
        "Write the complete OCI Point of View / PRFAQ for Northwind Health. Cast a "
        "visionary future press release, then answer both external customer questions "
        "and internal Oracle discovery questions. Treat all targets as proposed validation "
        "targets, not achieved outcomes."
    ),
    "jep": (
        "Create the complete Joint Execution Plan for the Northwind Health three-tier "
        ".NET member claims portal POC. Duration is 6 weeks in us-chicago-1. Phase 1 "
        "Assessment is Week 1; Phase 2 Build is Weeks 2-3; Phase 3 Validate is Weeks 4-6. "
        "In scope: OCI WAF, Flexible Load Balancer, two E5 Flex web servers, two E5 Flex "
        "application servers, Oracle Base Database Service, 500 GB Block Volume, 100 GB "
        "Object Storage, Logging, Monitoring, IAM, Vault, HIPAA control validation, and "
        "evidence collection. Out of scope: production cutover and production data migration. "
        "The architecture is public WAF and load balancer ingress to private web, application, "
        "and database tiers. Success criteria and evidence: measure p95 response time under "
        "300 ms, 99.9% availability during the validation window, recovery time under 60 "
        "minutes, and a proposed 30% infrastructure cost reduction target within 12 months. "
        "Owners and commitments: Oracle solution architect owns architecture and evidence; "
        "Northwind application owner runs tests; Northwind security lead reviews HIPAA controls. "
        "Track access delay, performance miss, and recovery miss risks with joint mitigation. "
        "Approvals are joint Oracle and Northwind go/no-go review. Include Bill of Materials, "
        "test cases, participants, deliverables, logistics, and handoff deliverables sections."
    ),
    "bom": "Generate the priced BOM for the fixed Northwind Health OCI architecture and sizing.",
    "diagram": (
        "Generate a draw.io architecture diagram for Northwind Health in us-chicago-1. "
        "Include an external user, OCI WAF, public Flexible Load Balancer, VCN, public and "
        "private subnets, two E5 Flex web servers, two E5 Flex application servers, "
        "Oracle Base Database Service, Block Volume, Object Storage, IAM, Vault, Logging, "
        "Monitoring, NSGs, route tables, Internet Gateway, NAT Gateway, and Service Gateway."
    ),
    "waf": (
        "Perform a complete OCI Well-Architected Framework review of the fixed Northwind "
        "Health architecture. Cover and score all six pillars, identify evidence gaps, "
        "map HIPAA-relevant controls, prioritize risks, and give actionable remediations."
    ),
    "terraform": (
        "Generate a complete Terraform bundle for the fixed Northwind Health architecture "
        "in us-chicago-1. Use variables and outputs, private application/database tiers, "
        "WAF and load-balancer ingress, NSGs, IAM, Vault, Logging, and Monitoring."
    ),
}


_EXPECTED_BOM_QUANTITIES = {
    "B97384": Decimal("12"),
    "B97385": Decimal("96"),
    "B91961": Decimal("500"),
    "B91962": Decimal("5000"),
    "B91628": Decimal("100"),
    "B99060": Decimal("4"),
}


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _slug(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", str(value).strip().lower()).strip("_")


def load_rubric(artifact_type: str, root: Path = RUBRIC_ROOT) -> dict[str, Any]:
    """Load one anchored rubric and return stable dimension identifiers."""
    path = root / f"{artifact_type}.md"
    text = path.read_text(encoding="utf-8")
    matches = list(
        re.finditer(
            r"^## Dimension:\s*(.+?)\s*$\n(.*?)(?=^## Dimension:|\Z)",
            text,
            flags=re.MULTILINE | re.DOTALL,
        )
    )
    dimensions: dict[str, dict[str, str]] = {}
    for match in matches:
        name = _slug(match.group(1))
        block = match.group(2).strip()
        high = re.search(r"^5\s*=\s*(.+)$", block, re.MULTILINE)
        low = re.search(r"^1\s*=\s*(.+)$", block, re.MULTILINE)
        if not high or not low:
            raise ValueError(f"Rubric dimension {name!r} lacks explicit 5/1 anchors")
        dimensions[name] = {
            "title": match.group(1).strip(),
            "five": high.group(1).strip(),
            "one": low.group(1).strip(),
            "text": block,
        }
    if not dimensions:
        raise ValueError(f"No dimensions found in {path}")
    return {"artifact_type": artifact_type, "path": str(path), "text": text, "dimensions": dimensions}


def canonical_sections(artifact_type: str, root: Path = GOLDEN_ROOT) -> list[str]:
    """Read canonical numbered sections from the SE-provided golden source note."""
    source = root / artifact_type / "SOURCE.md"
    if not source.exists():
        return []
    return [
        re.split(r"\s+[—–-]\s+", match.group(1).strip(), maxsplit=1)[0]
        for match in re.finditer(r"^\d+\.\s+(.+?)\s*$", source.read_text(encoding="utf-8"), re.MULTILINE)
    ]


def load_golden_exemplar(artifact_type: str, root: Path = GOLDEN_ROOT) -> dict[str, str] | None:
    """Load the first real exemplar, or return None for an intentionally empty type."""
    directory = root / artifact_type
    if not directory.exists():
        return None
    candidates = sorted(
        path for path in directory.iterdir()
        if path.is_file() and path.name not in {"SOURCE.md", ".gitkeep"}
    )
    if not candidates:
        return None
    path = candidates[0]
    if path.suffix.lower() == ".docx":
        content = _docx_text(path.read_bytes())
    elif path.suffix.lower() == ".xlsx":
        content = _xlsx_text(path.read_bytes())
    else:
        content = path.read_text(encoding="utf-8", errors="replace")
    try:
        # Record repo-relative so the baseline JSON resolves from any checkout.
        recorded = str(path.relative_to(ROOT))
    except ValueError:
        recorded = str(path)
    return {"path": recorded, "content": content}


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    return "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    ).strip()


def _xlsx_text(content: bytes) -> str:
    workbook = openpyxl.load_workbook(BytesIO(content), data_only=True, read_only=True)
    try:
        return "\n".join(
            " | ".join("" if value is None else str(value) for value in row)
            for sheet in workbook.worksheets
            for row in sheet.iter_rows(values_only=True)
            if any(value is not None for value in row)
        )
    finally:
        workbook.close()


def _check(passed: bool, detail: str) -> dict[str, Any]:
    return {"passed": bool(passed), "detail": str(detail)}


def objective_checks(
    artifact_type: str,
    content: str,
    fixture: dict[str, Any] = FIXTURE,
) -> dict[str, dict[str, Any]]:
    """Run deterministic type-specific structure, grounding, and format checks."""
    try:
        checks = {
            "pov": _check_pov,
            "jep": _check_jep,
            "bom": _check_bom,
            "diagram": _check_diagram,
            "waf": _check_waf,
            "terraform": _check_terraform,
        }[artifact_type](content, fixture)
    except Exception as exc:
        checks = {"validator_execution": _check(False, f"validator failed: {exc}")}
    return checks


def _headings(text: str) -> list[str]:
    return [match.group(1).strip() for match in re.finditer(r"^#{1,6}\s+(.+?)\s*$", text, re.MULTILINE)]


def _normalize_section(value: str) -> str:
    normalized = re.sub(r"^\d+[.)]?\s*", "", value.lower())
    normalized = re.sub(r"\bpoc\b", "proof of concept", normalized)
    return re.sub(r"[^a-z0-9]+", " ", normalized).strip()


def _missing_sections(text: str, artifact_type: str) -> list[str]:
    expected = canonical_sections(artifact_type)
    actual = [_normalize_section(item) for item in _headings(text)]
    return [
        section for section in expected
        if not any(_normalize_section(section) in heading or heading in _normalize_section(section) for heading in actual)
    ]


def _unsupported_quantified_values(text: str) -> list[str]:
    allowed = {
        ("30", "%"), ("300", "ms"), ("99.9", "%"), ("60", "minutes"),
        ("6", "weeks"), ("12", "months"), ("2", "ocpu"), ("4", "ocpu"),
        ("12", "ocpu"), ("16", "gb"), ("32", "gb"), ("64", "gb"),
        ("96", "gb"), ("500", "gb"), ("100", "gb"), ("730", "hours"),
    }
    aliases = {
        "millisecond": "ms", "milliseconds": "ms", "minute": "minutes",
        "hour": "hours", "week": "weeks", "month": "months", "ocpus": "ocpu",
    }
    found: list[str] = []
    pattern = re.compile(
        r"\b(\d+(?:\.\d+)?)\s*-?\s*(%|ms|milliseconds?|minutes?|hours?|weeks?|months?|gb|tb|ocpus?)\b",
        re.IGNORECASE,
    )
    for number, unit in pattern.findall(text):
        normalized_unit = aliases.get(unit.lower(), unit.lower())
        pair = (number.rstrip("0").rstrip(".") if "." in number else number, normalized_unit)
        if pair not in allowed:
            found.append(f"{number} {unit}")
    return list(dict.fromkeys(found))


def _check_pov(content: str, fixture: dict[str, Any]) -> dict[str, dict[str, Any]]:
    missing = _missing_sections(content, "pov")
    unsupported = _unsupported_quantified_values(content)
    lowered = content.lower()
    return {
        "markdown_parse": _check(bool(_headings(content)), f"headings={len(_headings(content))}"),
        "canonical_sections": _check(not missing, "missing=" + ", ".join(missing) if missing else "all golden sections present"),
        "customer_grounding": _check(fixture["customer_name"].lower() in lowered, fixture["customer_name"]),
        "fixture_fidelity": _check(
            all(marker in lowered for marker in ("us-chicago-1", "hipaa", "claims")),
            "requires region, HIPAA, and claims workload",
        ),
        "quantified_claim_grounding": _check(not unsupported, "unsupported=" + ", ".join(unsupported) if unsupported else "all quantified claims trace"),
    }


def _check_jep(content: str, fixture: dict[str, Any]) -> dict[str, dict[str, Any]]:
    from agent.jep_docx_renderer import render_jep_docx

    missing = _missing_sections(content, "jep")
    phase_hits = [bool(re.search(rf"\bPhase\s*{index}\b", content, re.IGNORECASE)) for index in (1, 2, 3)]
    owner_hits = sum(marker in content.lower() for marker in ("oracle solution architect", "application owner", "security lead"))
    unsupported = _unsupported_quantified_values(content)
    try:
        docx = render_jep_docx(content, customer_name=fixture["customer_name"])
        parsed = _docx_text(docx)
        docx_ok = bool(parsed)
        docx_detail = f"parsed {len(parsed)} characters"
    except Exception as exc:
        docx_ok = False
        docx_detail = str(exc)
    return {
        "canonical_sections": _check(not missing, "missing=" + ", ".join(missing) if missing else "all 10 golden sections present"),
        "three_phase_plan": _check(all(phase_hits), f"phase_hits={phase_hits}"),
        "owners_present": _check(owner_hits >= 2, f"owner anchors={owner_hits}/3"),
        "customer_grounding": _check(fixture["customer_name"].lower() in content.lower(), fixture["customer_name"]),
        "quantified_claim_grounding": _check(not unsupported, "unsupported=" + ", ".join(unsupported) if unsupported else "all quantified claims trace"),
        "docx_parse": _check(docx_ok, docx_detail),
    }


def _json_object(content: str) -> dict[str, Any]:
    cleaned = str(content or "").strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    data = json.loads(cleaned)
    if not isinstance(data, dict):
        raise ValueError("expected a JSON object")
    return data


def _bom_payload(content: str) -> dict[str, Any]:
    data = _json_object(content)
    payload = data.get("bom_payload", data)
    if not isinstance(payload, dict):
        raise ValueError("BOM payload is not an object")
    return payload


def _decimal(value: Any) -> Decimal | None:
    if value in (None, "", "TBD"):
        return None
    try:
        return Decimal(str(value).replace(",", ""))
    except (InvalidOperation, ValueError):
        return None


def _line_monthly(row: dict[str, Any]) -> Decimal | None:
    extended = _decimal(row.get("extended_price"))
    if extended is not None:
        return extended
    quantity = _decimal(row.get("quantity"))
    price = _decimal(row.get("unit_price"))
    if quantity is None or price is None:
        return None
    multiplier = Decimal("730") if "hour" in str(row.get("metric") or "").lower() else Decimal("1")
    return quantity * price * multiplier


def _check_bom(content: str, fixture: dict[str, Any]) -> dict[str, dict[str, Any]]:
    from agent.bom_service import get_shared_bom_service

    payload = _bom_payload(content)
    rows = [row for row in payload.get("line_items", []) if isinstance(row, dict)]
    required_fields = ("sku", "description", "quantity", "metric")
    complete = bool(rows) and all(all(row.get(field) not in (None, "") for field in required_fields) for row in rows)
    priced = [row for row in rows if _decimal(row.get("unit_price")) is not None]
    monthly_values = [_line_monthly(row) for row in rows]
    calculable = [value for value in monthly_values if value is not None]
    calculated = sum(calculable, Decimal("0"))
    stated = _decimal(payload.get("monthly_total"))
    if stated is None and isinstance(payload.get("totals"), dict):
        stated = _decimal(payload["totals"].get("estimated_monthly_cost"))
    math_ok = stated is not None and abs(stated - calculated) <= Decimal("0.02")
    by_sku: dict[str, Decimal] = defaultdict(lambda: Decimal("0"))
    for row in rows:
        quantity = _decimal(row.get("quantity"))
        if row.get("sku") and quantity is not None:
            by_sku[str(row["sku"])] += quantity
    mismatches = {
        sku: f"expected {expected}, got {by_sku.get(sku)}"
        for sku, expected in _EXPECTED_BOM_QUANTITIES.items()
        if by_sku.get(sku) != expected
    }
    try:
        workbook_bytes = get_shared_bom_service().generate_xlsx(payload)
        workbook = openpyxl.load_workbook(BytesIO(workbook_bytes), data_only=False, read_only=True)
        sheet_count = len(workbook.sheetnames)
        workbook.close()
        xlsx_ok = sheet_count > 0
        xlsx_detail = f"parsed {sheet_count} worksheet(s)"
    except Exception as exc:
        xlsx_ok = False
        xlsx_detail = str(exc)
    actual_services = {str(row.get("canonical_service_id") or "") for row in rows}
    expected_services = set(fixture["structured_inputs"]["target_service_ids"])
    missing_services = sorted(expected_services - actual_services)
    return {
        "line_item_structure": _check(complete, f"line_items={len(rows)}"),
        "priced_line_items": _check(len(priced) == len(rows), f"priced={len(priced)}/{len(rows)}"),
        "monthly_math": _check(math_ok, f"stated={stated} calculated={calculated}"),
        "fixture_quantities": _check(not mismatches, json.dumps(mismatches, sort_keys=True) if mismatches else "requested quantities match"),
        "service_grounding": _check(not missing_services, "missing=" + ", ".join(missing_services) if missing_services else "all requested services represented"),
        "region_grounding": _check(payload.get("region") == fixture["region"], f"region={payload.get('region')}"),
        "xlsx_parse": _check(xlsx_ok, xlsx_detail),
    }


def _check_diagram(content: str, fixture: dict[str, Any]) -> dict[str, dict[str, Any]]:
    try:
        root = ElementTree.fromstring(content)
        tag = root.tag.lower()
        parse_ok = "mxfile" in tag or "mxgraphmodel" in tag
        parse_detail = root.tag
    except Exception as exc:
        parse_ok = False
        parse_detail = str(exc)
    lowered = content.lower()
    tiers = {
        "web": any(marker in lowered for marker in ("web tier", "web server", "iis")),
        "application": any(
            marker in lowered
            for marker in (
                "application tier", "app tier", "claims app",
                "application subnet", "app subnet", "private application",
            )
        ),
        "database": any(marker in lowered for marker in ("database", "oracle base")),
    }
    required = ("waf", "load balancer", "object storage", "block volume", "logging", "monitoring")
    missing = [marker for marker in required if marker not in lowered]
    return {
        "drawio_parse": _check(parse_ok, parse_detail),
        "three_tier_structure": _check(all(tiers.values()), json.dumps(tiers, sort_keys=True)),
        "service_grounding": _check(not missing, "missing=" + ", ".join(missing) if missing else "required services represented"),
        "region_grounding": _check(fixture["region"] in lowered, fixture["region"]),
        "customer_grounding": _check(fixture["customer_name"].lower() in lowered, fixture["customer_name"]),
    }


def _check_waf(content: str, fixture: dict[str, Any]) -> dict[str, dict[str, Any]]:
    lowered = content.lower().replace("cost optimization", "cost optimisation")
    pillars = (
        "security", "reliability", "performance efficiency", "cost optimisation",
        "operational excellence", "continuous improvement",
    )
    missing = [pillar for pillar in pillars if pillar not in lowered]
    score_hits = 0
    for pillar in pillars:
        # Match from the heading line itself (not just the body after it) —
        # producers commonly put "Score: N/5" directly in the heading text,
        # e.g. "## Pillar 6: Sustainability (Score: 3/5 - ...)".
        section = re.search(
            rf"^##[^\n]*{re.escape(pillar)}[^\n]*\n.*?(?=^##\s|\Z)",
            lowered,
            flags=re.MULTILINE | re.DOTALL,
        )
        if section and re.search(
            r"\bscore\s*:\s*(?:[*_]{1,2})?[1-5](?:\.\d+)?\s*/\s*5",
            section.group(0),
        ):
            score_hits += 1
    return {
        "markdown_parse": _check(bool(_headings(content)), f"headings={len(_headings(content))}"),
        "six_pillars": _check(not missing, "missing=" + ", ".join(missing) if missing else "all pillars present"),
        "pillar_scores": _check(score_hits == 6, f"scored={score_hits}/6"),
        "customer_grounding": _check(fixture["customer_name"].lower() in lowered, fixture["customer_name"]),
        "architecture_grounding": _check(all(marker in lowered for marker in ("waf", "load balancer", "hipaa")), "requires WAF, load balancer, and HIPAA"),
    }


def _terraform_files(content: str) -> dict[str, str]:
    data = _json_object(content)
    files = data.get("files", data)
    if not isinstance(files, dict):
        raise ValueError("Terraform files payload is not an object")
    return {str(key): str(value or "") for key, value in files.items()}


def _terraform_validate(files: dict[str, str]) -> tuple[bool, str]:
    executable = shutil.which("terraform")
    if not executable:
        return False, "terraform CLI unavailable"
    filename_map = {
        "main_tf": "main.tf", "variables_tf": "variables.tf",
        "outputs_tf": "outputs.tf", "readme_md": "README.md",
    }
    with tempfile.TemporaryDirectory(prefix="subagent-quality-tf-") as directory:
        root = Path(directory)
        for key, filename in filename_map.items():
            if files.get(key):
                (root / filename).write_text(files[key], encoding="utf-8")
        init = subprocess.run(
            [executable, "init", "-backend=false", "-input=false", "-no-color"],
            cwd=root, capture_output=True, text=True, timeout=120,
        )
        if init.returncode != 0:
            return False, "terraform init failed: " + (init.stderr or init.stdout).strip()[-1000:]
        validate = subprocess.run(
            [executable, "validate", "-no-color"],
            cwd=root, capture_output=True, text=True, timeout=120,
        )
        detail = (validate.stdout if validate.returncode == 0 else validate.stderr or validate.stdout).strip()
        return validate.returncode == 0, detail[-1000:]


def _check_terraform(content: str, fixture: dict[str, Any]) -> dict[str, dict[str, Any]]:
    files = _terraform_files(content)
    required = ("main_tf", "variables_tf", "outputs_tf", "readme_md")
    missing = [key for key in required if not files.get(key, "").strip()]
    joined = "\n".join(files.values()).lower()
    validate_ok, validate_detail = _terraform_validate(files)
    resource_hits = sum(marker in joined for marker in ("oci_core_vcn", "oci_core_instance", "oci_load_balancer", "oci_database"))
    return {
        "bundle_structure": _check(not missing, "missing=" + ", ".join(missing) if missing else "all four files present"),
        "terraform_validate": _check(validate_ok, validate_detail),
        "resource_structure": _check(resource_hits >= 3, f"required resource-family hits={resource_hits}"),
        "region_grounding": _check(fixture["region"] in joined, fixture["region"]),
        "service_grounding": _check(all(marker in joined for marker in ("waf", "load", "vault", "log")), "requires WAF/load balancer/Vault/Logging"),
    }


def _render_rubric(rubric: dict[str, Any]) -> str:
    return "\n\n".join(
        f"{name}:\n5 = {definition['five']}\n1 = {definition['one']}"
        for name, definition in rubric["dimensions"].items()
    )


def _absolute_prompt(artifact_type: str, content: str, rubric: dict[str, Any]) -> str:
    dimensions = list(rubric["dimensions"])
    return f"""You are an independent artifact-quality judge. Score the candidate {artifact_type} only against the anchored rubric.

Return JSON only:
{{"scores": {{"dimension_key": {{"score": 1, "rationale": "specific evidence"}}}}}}
Use every exact dimension key: {json.dumps(dimensions)}. Scores must be integers 1-5.

[FIXED ENGAGEMENT FIXTURE]
{json.dumps(FIXTURE, ensure_ascii=False, sort_keys=True)}

[ANCHORED RUBRIC]
{_render_rubric(rubric)}

[CANDIDATE]
{content[:_JUDGE_CONTENT_LIMIT]}
"""


def _pairwise_prompt(
    artifact_type: str,
    candidate_a: str,
    candidate_b: str,
    rubric: dict[str, Any],
    *,
    labels: tuple[str, str] = ("a", "b"),
) -> str:
    dimensions = list(rubric["dimensions"])
    return f"""You are an independent artifact-quality judge. Compare two {artifact_type} artifacts dimension by dimension. Do not reward verbosity or ordering.

Return JSON only:
{{"dimensions": {{"dimension_key": {{"winner": "{labels[0]}|{labels[1]}|tie", "rationale": "specific comparative evidence"}}}}}}
Use every exact dimension key: {json.dumps(dimensions)}.

[FIXED ENGAGEMENT FIXTURE]
{json.dumps(FIXTURE, ensure_ascii=False, sort_keys=True)}

[ANCHORED RUBRIC]
{_render_rubric(rubric)}

[{labels[0].upper()}]
{candidate_a[:_JUDGE_CONTENT_LIMIT]}

[{labels[1].upper()}]
{candidate_b[:_JUDGE_CONTENT_LIMIT]}
"""


def _parse_json_response(raw: Any) -> dict[str, Any]:
    if isinstance(raw, dict):
        return raw
    text = str(raw or "").strip()
    fenced = re.search(r"```(?:json)?\s*(\{.*\})\s*```", text, re.DOTALL | re.IGNORECASE)
    if fenced:
        text = fenced.group(1)
    else:
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if match:
            text = match.group(0)
    data = json.loads(text)
    if not isinstance(data, dict):
        raise ValueError("judge response is not a JSON object")
    return data


async def _invoke_judge(
    judge_runner: Callable[..., Any],
    prompt: str,
    *,
    model_id: str,
    artifact_type: str,
    mode: str,
    dimensions: list[str],
) -> Any:
    kwargs = {
        "prompt": prompt,
        "model_id": model_id,
        "artifact_type": artifact_type,
        "mode": mode,
        "dimensions": dimensions,
    }
    if inspect.iscoroutinefunction(judge_runner):
        return await judge_runner(**kwargs)
    return await asyncio.to_thread(judge_runner, **kwargs)


async def _absolute_judgments(
    artifact_type: str,
    content: str,
    rubric: dict[str, Any],
    judge_runner: Callable[..., Any],
    model_id: str,
    runs: int,
) -> list[dict[str, Any]]:
    dimensions = list(rubric["dimensions"])
    prompt = _absolute_prompt(artifact_type, content, rubric)
    judgments: list[dict[str, Any]] = []
    for index in range(1, runs + 1):
        print(f"Judge {artifact_type} absolute {index}/{runs}", flush=True)
        try:
            raw = await _invoke_judge(
                judge_runner, prompt, model_id=model_id, artifact_type=artifact_type,
                mode="absolute", dimensions=dimensions,
            )
            parsed = _parse_json_response(raw)
            scores = parsed.get("scores") if isinstance(parsed.get("scores"), dict) else {}
            normalized: dict[str, Any] = {}
            for dimension in dimensions:
                item = scores.get(dimension) if isinstance(scores.get(dimension), dict) else {}
                score = int(item.get("score"))
                if score < 1 or score > 5:
                    raise ValueError(f"score outside 1-5 for {dimension}")
                normalized[dimension] = {"score": score, "rationale": str(item.get("rationale") or "")}
            judgments.append({"run": index, "status": "ok", "scores": normalized, "raw": raw})
        except Exception as exc:
            judgments.append({"run": index, "status": "error", "error": str(exc)})
    return judgments


async def _pairwise_judgments(
    artifact_type: str,
    content_a: str,
    content_b: str,
    rubric: dict[str, Any],
    judge_runner: Callable[..., Any],
    model_id: str,
    runs: int,
    *,
    labels: tuple[str, str] = ("a", "b"),
) -> list[dict[str, Any]]:
    dimensions = list(rubric["dimensions"])
    prompt = _pairwise_prompt(artifact_type, content_a, content_b, rubric, labels=labels)
    judgments: list[dict[str, Any]] = []
    allowed = {labels[0], labels[1], "tie"}
    for index in range(1, runs + 1):
        print(
            f"Judge {artifact_type} pairwise {labels[0]} vs {labels[1]} {index}/{runs}",
            flush=True,
        )
        try:
            raw = await _invoke_judge(
                judge_runner, prompt, model_id=model_id, artifact_type=artifact_type,
                mode="pairwise", dimensions=dimensions,
            )
            parsed = _parse_json_response(raw)
            values = parsed.get("dimensions") if isinstance(parsed.get("dimensions"), dict) else {}
            normalized: dict[str, Any] = {}
            for dimension in dimensions:
                item = values.get(dimension) if isinstance(values.get(dimension), dict) else {}
                winner = str(item.get("winner") or "").lower()
                if winner not in allowed:
                    raise ValueError(f"invalid pairwise winner for {dimension}: {winner}")
                normalized[dimension] = {"winner": winner, "rationale": str(item.get("rationale") or "")}
            judgments.append({"run": index, "status": "ok", "dimensions": normalized, "raw": raw})
        except Exception as exc:
            judgments.append({"run": index, "status": "error", "error": str(exc)})
    return judgments


def _score_distribution(artifacts: list[dict[str, Any]], dimensions: list[str]) -> dict[str, Any]:
    result: dict[str, Any] = {}
    for dimension in dimensions:
        values = [
            int(judgment["scores"][dimension]["score"])
            for artifact in artifacts
            for judgment in artifact.get("judgments", [])
            if judgment.get("status") == "ok" and dimension in judgment.get("scores", {})
        ]
        result[dimension] = {
            "values": values,
            "count": len(values),
            "min": min(values) if values else None,
            "median": statistics.median(values) if values else None,
            "max": max(values) if values else None,
        }
    return result


def _objective_distribution(artifacts: list[dict[str, Any]]) -> dict[str, Any]:
    names = sorted({name for artifact in artifacts for name in artifact.get("objective", {})})
    checks: dict[str, Any] = {}
    for name in names:
        values = [artifact["objective"][name]["passed"] for artifact in artifacts if name in artifact.get("objective", {})]
        checks[name] = {"passed": sum(values), "total": len(values), "pass_rate": round(sum(values) / len(values), 4) if values else None}
    overall = [artifact.get("objective_pass", False) for artifact in artifacts]
    return {
        "artifact_passes": sum(overall),
        "artifact_total": len(overall),
        "artifact_pass_rate": round(sum(overall) / len(overall), 4) if overall else None,
        "checks": checks,
    }


def _pairwise_distribution(
    artifacts: list[dict[str, Any]], dimensions: list[str], field: str = "golden_judgments"
) -> dict[str, Any]:
    result: dict[str, Any] = {}
    for dimension in dimensions:
        winners = [
            judgment["dimensions"][dimension]["winner"]
            for artifact in artifacts
            for judgment in artifact.get(field, [])
            if judgment.get("status") == "ok" and dimension in judgment.get("dimensions", {})
        ]
        counts = {key: winners.count(key) for key in ("candidate", "golden", "a", "b", "tie") if winners.count(key)}
        result[dimension] = {
            "count": len(winners),
            "wins": counts,
            "candidate_win_rate": round(winners.count("candidate") / len(winners), 4) if winners else None,
            "a_win_rate": round(winners.count("a") / len(winners), 4) if winners else None,
            "b_win_rate": round(winners.count("b") / len(winners), 4) if winners else None,
        }
    return result


async def run_baseline(
    *,
    producer_runs: int,
    judge_runs: int,
    artifact_types: tuple[str, ...],
    judge_model_id: str,
    producer: Callable[..., Any],
    card_loader: Callable[..., Any],
    judge_runner: Callable[..., Any],
) -> dict[str, Any]:
    """Generate and score the fixed fixture through the existing A2A boundary."""
    report: dict[str, Any] = {
        "schema_version": "1.0",
        "mode": "baseline",
        "started_at": _now(),
        "fixture": FIXTURE,
        "producer_runs": producer_runs,
        "judge_runs": judge_runs,
        "judge_model_id": judge_model_id,
        "environment": {
            "producer_source_commit": subprocess.run(
                ["git", "rev-parse", "--short", "HEAD"],
                cwd=ROOT,
                capture_output=True,
                text=True,
                check=False,
            ).stdout.strip(),
            "producer_transport": "isolated current-source A2A services",
            "judge_model_distinct_from_producers": True,
            "terraform_cli_available": shutil.which("terraform") is not None,
        },
        "types": {},
    }
    for artifact_type in artifact_types:
        rubric = load_rubric(artifact_type)
        dimensions = list(rubric["dimensions"])
        card = card_loader(artifact_type)
        if inspect.isawaitable(card):
            card = await card
        producer_model_id = str((card or {}).get("llm_model_id") or "")
        if producer_model_id and producer_model_id == judge_model_id:
            raise ValueError(
                f"judge model must differ from {artifact_type} producer model {producer_model_id}"
            )
        golden = load_golden_exemplar(artifact_type)
        artifacts: list[dict[str, Any]] = []
        for run_index in range(1, producer_runs + 1):
            print(
                f"Generate {artifact_type} artifact {run_index}/{producer_runs}",
                flush=True,
            )
            trace_id = f"quality-{artifact_type}-{run_index}-{uuid.uuid4().hex[:8]}"
            response = producer(
                artifact_type,
                TASKS[artifact_type],
                engagement_context=FIXTURE,
                trace_id=trace_id,
            )
            if inspect.isawaitable(response):
                response = await response
            status = str((response or {}).get("status") or "error")
            content = str((response or {}).get("result") or "")
            objective = objective_checks(artifact_type, content) if status == "ok" else {
                "producer_status": _check(False, f"producer returned {status}")
            }
            artifact: dict[str, Any] = {
                "artifact_id": f"{artifact_type}-{run_index}",
                "run": run_index,
                "producer_model_id": producer_model_id,
                "producer_status": status,
                "trace": dict((response or {}).get("trace") or {}),
                "content": content,
                "objective": objective,
                "objective_pass": bool(objective) and all(item["passed"] for item in objective.values()),
                "judgments": [],
                "golden_judgments": [],
            }
            if status == "ok" and content:
                artifact["judgments"] = await _absolute_judgments(
                    artifact_type, content, rubric, judge_runner, judge_model_id, judge_runs
                )
                if golden is not None:
                    artifact["golden_judgments"] = await _pairwise_judgments(
                        artifact_type, content, golden["content"], rubric, judge_runner,
                        judge_model_id, judge_runs, labels=("candidate", "golden"),
                    )
            artifacts.append(artifact)
        report["types"][artifact_type] = {
            "rubric_path": rubric["path"],
            "dimensions": dimensions,
            "producer_model_id": producer_model_id,
            "golden": {"available": golden is not None, "path": golden["path"] if golden else ""},
            "objective": _objective_distribution(artifacts),
            "subjective": {"dimensions": _score_distribution(artifacts, dimensions)},
            "golden_pairwise": {
                "skipped": golden is None,
                "dimensions": _pairwise_distribution(artifacts, dimensions),
            },
            "artifacts": artifacts,
        }
    report["completed_at"] = _now()
    return report


def _artifacts_from_report(report: dict[str, Any], artifact_type: str) -> list[dict[str, Any]]:
    data = (report.get("types") or {}).get(artifact_type, {})
    artifacts = data.get("artifacts") if isinstance(data, dict) else []
    return [item for item in artifacts or [] if isinstance(item, dict) and item.get("content")]


async def run_pairwise(
    report_a: dict[str, Any],
    report_b: dict[str, Any],
    *,
    judge_runs: int,
    judge_model_id: str,
    judge_runner: Callable[..., Any],
) -> dict[str, Any]:
    producer_models = {
        str(artifact.get("producer_model_id") or "")
        for artifact_type in ARTIFACT_TYPES
        for report in (report_a, report_b)
        for artifact in _artifacts_from_report(report, artifact_type)
        if artifact.get("producer_model_id")
    }
    if judge_model_id in producer_models:
        raise ValueError(
            f"judge model must differ from compared producer model {judge_model_id}"
        )
    result: dict[str, Any] = {
        "schema_version": "1.0", "mode": "pairwise", "created_at": _now(),
        "judge_model_id": judge_model_id, "judge_runs": judge_runs, "types": {},
    }
    for artifact_type in ARTIFACT_TYPES:
        artifacts_a = _artifacts_from_report(report_a, artifact_type)
        artifacts_b = _artifacts_from_report(report_b, artifact_type)
        pairs = min(len(artifacts_a), len(artifacts_b))
        if pairs == 0:
            result["types"][artifact_type] = {"pairs": 0, "skipped": True, "dimensions": {}, "raw_judgments": []}
            continue
        rubric = load_rubric(artifact_type)
        dimensions = list(rubric["dimensions"])
        raw: list[dict[str, Any]] = []
        for index in range(pairs):
            judgments = await _pairwise_judgments(
                artifact_type,
                str(artifacts_a[index]["content"]),
                str(artifacts_b[index]["content"]),
                rubric,
                judge_runner,
                judge_model_id,
                judge_runs,
                labels=("a", "b"),
            )
            raw.extend({**item, "pair": index + 1} for item in judgments)
        holder = [{"pairwise": raw}]
        result["types"][artifact_type] = {
            "pairs": pairs,
            "skipped": False,
            "dimensions": _pairwise_distribution(holder, dimensions, field="pairwise"),
            "raw_judgments": raw,
        }
    return result


def _pearson(xs: list[float], ys: list[float]) -> float | None:
    if len(xs) < 2 or len(xs) != len(ys):
        return None
    mean_x = statistics.mean(xs)
    mean_y = statistics.mean(ys)
    numerator = sum((x - mean_x) * (y - mean_y) for x, y in zip(xs, ys))
    denominator = math.sqrt(
        sum((x - mean_x) ** 2 for x in xs) * sum((y - mean_y) ** 2 for y in ys)
    )
    return round(numerator / denominator, 6) if denominator else None


async def run_calibration(
    payload: dict[str, Any],
    *,
    judge_runs: int,
    judge_model_id: str,
    judge_runner: Callable[..., Any] | None,
) -> dict[str, Any]:
    """Correlate judge and SE scores, optionally judging labelled artifact content."""
    records = [dict(item) for item in payload.get("records", []) if isinstance(item, dict)]
    for label_index, label in enumerate(payload.get("labels", []) or [], start=1):
        if not isinstance(label, dict):
            continue
        artifact_type = str(label.get("artifact_type") or "")
        content = str(label.get("content") or label.get("artifact") or "")
        se_scores = label.get("se_scores") if isinstance(label.get("se_scores"), dict) else {}
        if not artifact_type or not content or judge_runner is None:
            continue
        rubric = load_rubric(artifact_type)
        judgments = await _absolute_judgments(
            artifact_type, content, rubric, judge_runner, judge_model_id, judge_runs
        )
        for dimension, se_score in se_scores.items():
            values = [
                item["scores"][dimension]["score"]
                for item in judgments
                if item.get("status") == "ok" and dimension in item.get("scores", {})
            ]
            if values:
                records.append({
                    "label": label.get("id") or f"label-{label_index}",
                    "artifact_type": artifact_type,
                    "dimension": dimension,
                    "se_score": float(se_score),
                    "judge_score": statistics.mean(values),
                })
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for record in records:
        if record.get("dimension") and record.get("se_score") is not None and record.get("judge_score") is not None:
            grouped[str(record["dimension"])].append(record)
    dimensions: dict[str, Any] = {}
    for dimension, items in sorted(grouped.items()):
        se_values = [float(item["se_score"]) for item in items]
        judge_values = [float(item["judge_score"]) for item in items]
        dimensions[dimension] = {
            "count": len(items),
            "pearson_correlation": _pearson(se_values, judge_values),
            "mean_absolute_error": round(statistics.mean(abs(a - b) for a, b in zip(se_values, judge_values)), 6),
        }
    return {
        "schema_version": "1.0", "mode": "calibration", "created_at": _now(),
        "judge_model_id": judge_model_id, "dimensions": dimensions, "records": records,
    }


class RealJudge:
    """OCI inference adapter for the configured independent judge model."""

    def __init__(self, model_id: str = "", min_interval: float = 0.0) -> None:
        from agent.runtime_config import resolve_agent_llm_config

        config = yaml.safe_load((ROOT / "config.yaml").read_text(encoding="utf-8")) or {}
        llm = resolve_agent_llm_config(config, "native_orchestrator")
        self.model_id = str(model_id or llm.get("model_id") or "")
        self.endpoint = str(llm.get("service_endpoint") or "")
        self.compartment_id = str(config.get("compartment_id") or "")
        self.max_tokens = int(llm.get("max_tokens") or 4000)
        self.top_p = float(llm.get("top_p") or 0.9)
        self.min_interval = max(0.0, float(min_interval))
        self._last_started = 0.0

    def __call__(self, *, prompt: str, model_id: str, **_kwargs: Any) -> str:
        from agent.llm_inference_client import run_inference

        remaining = self.min_interval - (time.monotonic() - self._last_started)
        if remaining > 0:
            time.sleep(remaining)
        self._last_started = time.monotonic()
        return run_inference(
            prompt=prompt,
            system_message=(
                "You are an independent quality judge. Follow the anchored rubric, "
                "cite artifact evidence, and return only valid JSON."
            ),
            endpoint=self.endpoint,
            model_id=str(model_id or self.model_id),
            compartment_id=self.compartment_id,
            max_tokens=self.max_tokens,
            temperature=0.2,
            top_p=self.top_p,
            top_k=0,
        )


async def _default_producer(
    artifact_type: str,
    task: str,
    *,
    engagement_context: dict[str, Any],
    trace_id: str,
) -> dict[str, Any]:
    from agent.sub_agent_client import call_sub_agent

    return await call_sub_agent(
        artifact_type, task, engagement_context=engagement_context, trace_id=trace_id
    )


async def _default_card_loader(artifact_type: str) -> dict[str, Any]:
    from agent.sub_agent_client import get_agent_card

    return await get_agent_card(artifact_type)


def _write_report(path: Path, report: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def _parse_types(value: str) -> tuple[str, ...]:
    if not value or value.lower() == "all":
        return ARTIFACT_TYPES
    selected = tuple(dict.fromkeys(item.strip().lower() for item in value.split(",") if item.strip()))
    unknown = sorted(set(selected) - set(ARTIFACT_TYPES))
    if unknown:
        raise ValueError("unknown artifact types: " + ", ".join(unknown))
    return selected


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    mode = parser.add_mutually_exclusive_group()
    mode.add_argument("--pairwise", nargs=2, metavar=("A_JSON", "B_JSON"))
    mode.add_argument("--calibrate", metavar="SE_SCORES_JSON")
    parser.add_argument("--runs", type=int, default=1, help="Producer generations per artifact type.")
    parser.add_argument("--judge-runs", type=int, default=3, help="Judge samples per artifact or pair.")
    parser.add_argument("--judge-model-id", default="", help="Independent judge model OCID.")
    parser.add_argument(
        "--min-judge-interval",
        type=float,
        default=0.0,
        help="Minimum seconds between live judge calls for service quota control.",
    )
    parser.add_argument("--types", default="all", help="Comma-separated artifact types or all.")
    parser.add_argument("--output", type=Path, default=None)
    return parser.parse_args()


async def _run_cli(args: argparse.Namespace) -> tuple[dict[str, Any], Path]:
    if args.runs < 1 or args.judge_runs < 1:
        raise ValueError("--runs and --judge-runs must be at least 1")
    judge = RealJudge(args.judge_model_id, args.min_judge_interval)
    judge_model_id = judge.model_id
    if not judge_model_id:
        raise ValueError("A judge model_id is required")
    if args.pairwise:
        report_a = json.loads(Path(args.pairwise[0]).read_text(encoding="utf-8"))
        report_b = json.loads(Path(args.pairwise[1]).read_text(encoding="utf-8"))
        report = await run_pairwise(
            report_a, report_b, judge_runs=args.judge_runs,
            judge_model_id=judge_model_id, judge_runner=judge,
        )
        output = args.output or ROOT / "docs" / "subagent-quality-pairwise.json"
    elif args.calibrate:
        payload = json.loads(Path(args.calibrate).read_text(encoding="utf-8"))
        report = await run_calibration(
            payload, judge_runs=args.judge_runs,
            judge_model_id=judge_model_id, judge_runner=judge,
        )
        output = args.output or ROOT / "docs" / "subagent-quality-calibration.json"
    else:
        report = await run_baseline(
            producer_runs=args.runs,
            judge_runs=args.judge_runs,
            artifact_types=_parse_types(args.types),
            judge_model_id=judge_model_id,
            producer=_default_producer,
            card_loader=_default_card_loader,
            judge_runner=judge,
        )
        output = args.output or DEFAULT_REPORT
    _write_report(output, report)
    return report, output


def main() -> int:
    try:
        report, output = asyncio.run(_run_cli(_parse_args()))
    except Exception as exc:
        print(f"quality evaluation failed: {exc}")
        return 1
    print(f"Wrote {report['mode']} quality report: {output}")
    if report["mode"] == "baseline":
        for artifact_type, result in report["types"].items():
            objective = result["objective"]
            print(
                f"  {artifact_type}: objective {objective['artifact_passes']}/"
                f"{objective['artifact_total']}; judge dimensions="
                f"{len(result['subjective']['dimensions'])}"
            )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
