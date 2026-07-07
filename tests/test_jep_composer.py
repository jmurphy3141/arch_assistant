from __future__ import annotations

from io import BytesIO

from docx import Document

from agent.jep_composer import (
    CORE_SECTIONS,
    compose_jep,
    extract_jep_brief,
    validate_jep_markdown,
)
from agent.jep_docx_renderer import render_jep_docx


QUALIFIED = (
    "Create the final 14-day JEP and POC plan for Apex Retail to validate migration of an "
    "on-premises three-tier retail web application to OCI us-ashburn-1. Scope: WAF, public "
    "Flexible Load Balancer, two private VM.Standard.E5.Flex web servers, private PostgreSQL "
    "database, Object Storage, Block Volume, logging, and monitoring. Use exactly three phases: "
    "Phase 1 Assessment on days 1-3, Phase 2 Build on days 4-9, and Phase 3 Validate on days 10-14. "
    "Success criteria: 99.9% availability during a 48-hour soak test, p95 response time under "
    "500 milliseconds at 100 requests per second, and database restore within 60 minutes. "
    "Oracle SA and Apex Retail technical lead each commit 8 hours per week. Include at least "
    "three risks, a go/no-go sign-off with fallback, explicit out-of-scope items, a BOM section, "
    "timeline, owners, approvals, and handoff deliverables. Generate only the JEP artifact; do "
    "not generate a separate BOM workbook."
)


def test_extracts_validated_brief() -> None:
    brief, missing = extract_jep_brief(QUALIFIED, {})
    assert missing == []
    assert brief is not None
    assert brief.customer == "Apex Retail"
    assert brief.region == "us-ashburn-1"
    assert [phase.name for phase in brief.phases] == ["Assessment", "Build", "Validate"]
    assert len(brief.criteria) == 3
    assert len(brief.risks) == 3


def test_extracts_three_named_owners_with_shared_commitment() -> None:
    request = QUALIFIED.replace(
        "Oracle SA and Apex Retail technical lead each commit 8 hours per week.",
        "Oracle SA, Apex Retail application lead, and Apex Retail network engineer each commit 8 hours per week.",
    )
    brief, missing = extract_jep_brief(request, {})

    assert missing == []
    assert brief is not None
    assert [owner.role for owner in brief.owners] == [
        "Oracle SA",
        "Apex Retail application lead",
        "Apex Retail network engineer",
    ]
    assert all(owner.commitment == "8 hours per week" for owner in brief.owners)


def test_preserves_realistic_order_rate_and_concurrent_user_criteria() -> None:
    order_request = QUALIFIED.replace(
        "99.9% availability during a 48-hour soak test",
        "sustain 750 orders per minute",
    )
    order_brief, order_missing = extract_jep_brief(order_request, {})
    assert order_missing == []
    assert order_brief is not None
    assert "sustain 750 orders per minute" in order_brief.criteria

    user_request = QUALIFIED.replace(
        "99.9% availability during a 48-hour soak test",
        "support 1200 concurrent enrollment users",
    )
    user_brief, user_missing = extract_jep_brief(user_request, {})
    assert user_missing == []
    assert user_brief is not None
    assert "support 1200 concurrent enrollment users" in user_brief.criteria

    qualified_rate_request = QUALIFIED.replace(
        "99.9% availability during a 48-hour soak test",
        "process 250 supplier orders per minute",
    )
    qualified_brief, qualified_missing = extract_jep_brief(qualified_rate_request, {})
    assert qualified_missing == []
    assert qualified_brief is not None
    assert "process 250 supplier orders per minute" in qualified_brief.criteria


def test_incomplete_request_returns_targeted_questions_without_artifact() -> None:
    result = compose_jep("Create a JEP for ACME", {"customer_name": "ACME"})
    assert result.status == "needs_input"
    assert result.markdown == ""
    assert result.missing_fields == ("workload",)
    assert any("workload" in question for question in result.questions)


def test_general_jep_drafts_missing_criterion_as_tbd() -> None:
    result = compose_jep(
        QUALIFIED.replace(
            ", and database restore within 60 minutes",
            "",
        )
    )
    assert result.status == "ok"
    assert "[TBD]" in result.markdown


def test_customer_and_workload_produce_draft_with_tbd_logistics() -> None:
    result = compose_jep(
        "Draft a JEP for Northwind Health, which runs a .NET member portal.",
        {"customer_name": "Northwind Health"},
    )

    assert result.status == "ok"
    assert result.brief is not None
    assert result.brief.duration == "[TBD]"
    assert all(phase.window == "[TBD]" for phase in result.brief.phases)
    assert result.brief.owners[0].role == "[TBD]"
    assert result.brief.criteria == ("[TBD]", "[TBD]", "[TBD]")
    assert "Northwind Health" in result.markdown
    assert "Jordan Kim" not in result.markdown
    assert "20 days" not in result.markdown


def test_canonical_ten_section_order_and_grounded_execution_content() -> None:
    result = compose_jep(QUALIFIED)
    assert result.status == "ok"
    headings = [line[3:] for line in result.markdown.splitlines() if line.startswith("## ")]
    assert headings == list(CORE_SECTIONS)
    assert headings[-1] == "Logistics"
    assert "## Proof of Concept Test Cases" in result.markdown
    assert "## Bill of Materials" in result.markdown
    assert "## POC Participants" in result.markdown
    assert "## Deliverables" in result.markdown
    assert "This section does not create or authorize a separate BOM workbook." in result.markdown
    assert "Go/no-go approval:" in result.markdown
    assert validate_jep_markdown(result.markdown, result.brief) == []


def test_exact_same_engagement_artifact_references_are_preserved() -> None:
    result = compose_jep(
        QUALIFIED,
        {
            "artifact_context": {
                "diagram": {"diagram_key": "diagram/apex/v4.drawio"},
                "bom": {"xlsx_artifact_key": "bom/apex/v2.xlsx", "line_item_count": 8},
            }
        },
    )
    assert result.status == "ok"
    assert "Diagram: diagram/apex/v4.drawio" in result.markdown
    assert "BOM: bom/apex/v2.xlsx" in result.markdown
    assert "Line Item Count: 8" in result.markdown


def test_revision_changes_only_explicit_grounded_value() -> None:
    original = compose_jep(QUALIFIED).markdown
    revised = compose_jep(
        "Replace 8 hours per week with 10 hours per week.",
        {"prior_version": original, "feedback": "Replace 8 hours per week with 10 hours per week."},
    )
    assert revised.status == "ok"
    assert revised.markdown.replace("10 hours per week", "8 hours per week") == original


def test_render_contains_no_unscoped_provisioning_examples() -> None:
    markdown = compose_jep(QUALIFIED).markdown
    assert "FastConnect" not in markdown
    assert "ADB Dedicated" not in markdown
    assert "1–2 hours" not in markdown


def test_validator_rejects_extra_phase_and_unsupported_numeric_fact() -> None:
    result = compose_jep(QUALIFIED)
    extra_phase = result.markdown.replace(
        "| Phase 3 Validate |",
        "| Phase 4 Deploy | 24 hours | Unsupported | Unsupported |\n| Phase 3 Validate |",
    )
    findings = validate_jep_markdown(extra_phase, result.brief)
    assert any("exactly Phase 1" in finding for finding in findings)
    assert any("unsupported numeric facts" in finding for finding in findings)


def test_markdown_and_docx_contain_matching_grounded_content() -> None:
    result = compose_jep(QUALIFIED)
    payload = render_jep_docx(result.markdown, customer_name="Apex Retail")
    doc = Document(BytesIO(payload))
    text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    for value in ("Apex Retail", "us-ashburn-1", "99.9%", "500 milliseconds", "60 minutes"):
        assert value in result.markdown
        assert value in text


def test_jep_requires_and_references_selected_poc_and_finalized_bom() -> None:
    request = QUALIFIED + " Ground this JEP in the selected POC and finalized BOM."
    missing = compose_jep(request, {})
    assert missing.status == "ok"
    assert "[TBD]" in missing.markdown

    result = compose_jep(
        request,
        {
            "artifact_context": {
                "poc": {
                    "selected_option_name": "Apex Retail Core Workload Validation POC",
                    "artifact_key": "poc_plan/apex/v1.json",
                    "oci_services": ["PostgreSQL DB System", "Site-to-Site VPN"],
                },
                "bom": {
                    "xlsx_artifact_key": "bom/apex/final.xlsx",
                    "xlsx_filename": "final.xlsx",
                    "summary": "Confirmed POC scope",
                },
            }
        },
    )
    assert result.status == "ok"
    assert "Selected POC: Apex Retail Core Workload Validation POC (poc_plan/apex/v1.json)" in result.markdown
    assert "BOM: bom/apex/final.xlsx" in result.markdown
    assert "Site-to-Site VPN" in result.markdown


def test_selected_poc_jep_accepts_two_exact_criteria_and_natural_owner_role() -> None:
    request = (
        "Generate only the JEP artifact for HarborStone's selected POC in us-ashburn-1. "
        "Use a 15-day duration with Phase 1 Assessment on days 1-3, Phase 2 Build on "
        "days 4-10, and Phase 3 Validate on days 11-15. Success criteria: p95 under "
        "500 milliseconds, restore within 60 minutes. Oracle SA and HarborStone lead "
        "each commit 10 hours per week. Include at least three risks and a go/no-go "
        "sign-off with fallback."
    )
    result = compose_jep(request, {
        "customer_name": "HarborStone",
        "artifact_context": {
            "poc": {
                "selected_option_name": "HarborStone Core Workload Validation POC",
                "artifact_key": "poc_plan/harbor/v1.json",
                "oci_services": ["VM.Standard.E5.Flex", "PostgreSQL DB System"],
                "grounding": {
                    "success_criteria": [
                        "p95 under 500 milliseconds",
                        "restore within 60 minutes",
                    ]
                },
            },
            "bom": {"xlsx_artifact_key": "bom/harbor/final.xlsx"},
            "diagram": {"diagram_key": "diagram/harbor/final.drawio"},
        },
    })

    assert result.status == "ok"
    assert "p95 under 500 milliseconds" in result.markdown
    assert "restore within 60 minutes" in result.markdown
    assert "HarborStone lead" in result.markdown
    assert "diagram/harbor/final.drawio" in result.markdown


def test_selected_poc_jep_splits_process_and_recover_criteria() -> None:
    request = (
        "Generate only the JEP artifact for NovaGrid's selected POC in us-phoenix-1. "
        "Use a 15-day duration with Phase 1 Assessment on days 1-3, Phase 2 Build on "
        "days 4-10, and Phase 3 Validate on days 11-15. Success criteria: process 300 "
        "requests per second, recover within 30 minutes. Oracle SA and NovaGrid engineer "
        "each commit 8 hours per week. Include at least three risks and a go/no-go sign-off "
        "with fallback."
    )
    result = compose_jep(request, {
        "customer_name": "NovaGrid",
        "artifact_context": {
            "poc": {
                "selected_option_name": "NovaGrid Performance Validation POC",
                "oci_services": ["VM.Standard.E5.Flex"],
                "grounding": {
                    "success_criteria": [
                        "process 300 requests per second",
                        "recover within 30 minutes",
                    ]
                },
            },
            "bom": {"xlsx_artifact_key": "bom/nova.xlsx"},
            "diagram": {"diagram_key": "diagram/nova.drawio"},
        },
    })

    assert result.status == "ok"
    assert "process 300 requests per second" in result.markdown
    assert "recover within 30 minutes" in result.markdown
