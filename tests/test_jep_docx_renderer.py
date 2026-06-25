from io import BytesIO

import pytest
from docx import Document

import agent.jep_agent as jep_agent
from agent.document_store import (
    JEP_DOCX_CONTENT_TYPE,
    get_jep_docx,
    save_doc,
    save_jep_docx,
)
from agent.jep_docx_renderer import TEMPLATE_PATH, render_jep_docx
from agent.persistence_objectstore import InMemoryObjectStore
from server.services.bom_artifacts import build_artifact_manifest


def _doc_text(doc: Document) -> str:
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_c3e_jep_template_is_sanitized() -> None:
    doc = Document(str(TEMPLATE_PATH))
    text = _doc_text(doc)

    assert "Document Body Placeholder" in text
    assert "Enterprise demand for high-performance computing resources" not in text


def test_render_jep_docx_from_markdown() -> None:
    payload = render_jep_docx(
        """# Joint Execution Plan - ACME
*Confidential - Oracle Restricted*

## Overview
ACME will validate OCI for a production POC.

- Confirm tenancy quota
- Deploy OKE

## Success Criteria
1. Provision platform in under 2 hours
2. Meet P99 latency target

| Owner | Role |
|-------|------|
| Oracle | Solutions Architect |
| ACME | Technical Champion |
""",
        customer_name="ACME",
    )

    assert payload.startswith(b"PK")
    doc = Document(BytesIO(payload))
    text = _doc_text(doc)
    assert "Proof of Concept Joint Execution Plan (JEP)" in text
    assert "Customer name: ACME" in text
    assert "Document Body Placeholder" not in text
    assert "Section Heading Placeholder" not in text
    assert "Joint Execution Plan - ACME" in text
    assert "Confirm tenancy quota" in text
    assert "Oracle" in text
    assert "Technical Champion" in text
    assert doc.tables[0].style.name == Document(str(TEMPLATE_PATH)).tables[0].style.name


def test_jep_agent_prompt_uses_c3e_section_contract() -> None:
    prompt = jep_agent._PROMPT_TEMPLATE.format(
        customer_name="ACME",
        context_summary="",
        new_notes_section="Notes",
        feedback_section="",
        qa_section="",
        previous_jep_section="",
        bom_md="BOM context",
        diagram_ref="Diagram key: agent3/acme/poc/LATEST.json",
        duration="2 weeks",
    )

    expected_sections = [
        "## Executive Summary",
        "## Objectives",
        "## Scope",
        "## POC Architecture",
        "## Phased Execution Plan",
        "## Success Criteria",
        "## Resource Plan",
        "## Risk Registry",
        "## Approvals",
    ]
    for section in expected_sections:
        assert section in prompt
    assert [prompt.index(section) for section in expected_sections] == sorted(
        prompt.index(section) for section in expected_sections
    )
    assert "## High Level Scope and Approach" not in prompt
    assert "## POC Participants" not in prompt
    assert "## Logistics" not in prompt
    assert "Oracle Corporation | 2300 Oracle Way" not in prompt
    assert prompt.rstrip().endswith("| [TBD] | ACME | Customer Technical Lead |  | [TBD] |")


def test_render_jep_docx_missing_template_fails(tmp_path) -> None:
    missing = tmp_path / "missing.docx"

    with pytest.raises(FileNotFoundError):
        render_jep_docx("# JEP", template_path=missing)


def test_save_and_get_jep_docx() -> None:
    store = InMemoryObjectStore()
    saved_md = save_doc(store, "jep", "acme", "# JEP", {"source": "test"})
    docx_bytes = render_jep_docx("# Joint Execution Plan - ACME", customer_name="ACME")

    saved_docx = save_jep_docx(
        store,
        "acme",
        saved_md["version"],
        docx_bytes,
        {"source": "test"},
    )

    assert saved_docx["docx_key"] == "jep/acme/v1.docx"
    assert saved_docx["docx_filename"] == "v1.docx"
    assert store.head("customers/acme/jep/v1.docx")
    assert store.head("customers/acme/jep/LATEST.docx")
    assert get_jep_docx(store, "acme", "v1.docx") == docx_bytes

    version_data = store.get("customers/acme/jep/v1.docx")
    assert version_data == docx_bytes
    assert JEP_DOCX_CONTENT_TYPE


def test_artifact_manifest_includes_jep_docx_download() -> None:
    manifest = build_artifact_manifest(
        "acme",
        {
            "tool_calls": [
                {
                    "tool": "generate_jep",
                    "result_data": {
                        "docx_key": "jep/acme/v1.docx",
                        "docx_filename": "v1.docx",
                    },
                }
            ]
        },
    )

    assert manifest["downloads"] == [
        {
            "type": "jep",
            "tool": "generate_jep",
            "key": "jep/acme/v1.docx",
            "filename": "v1.docx",
            "download_url": "/api/jep/acme/download/v1.docx",
        }
    ]
